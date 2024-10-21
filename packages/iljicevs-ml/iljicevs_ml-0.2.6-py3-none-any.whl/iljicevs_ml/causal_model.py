import os
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from causalml.inference.meta import LRSRegressor
import shap
import pandas as pd
from openpyxl import load_workbook
from openpyxl.drawing.image import Image
from docx import Document
from docx.shared import Inches

class IljicevsCausalModel:
    """
    Class for working with a causal model (EN).
    Класс для работы с причинно-следственной моделью (RU).
    Klasse zur Arbeit mit einem Kausalmodell (DE).
    用于处理因果模型的类 (ZH).
    """
    def __init__(self, model_type='meta'):
        """
        Initialize the causal model (EN).
        Инициализация причинно-следственной модели (RU).
        Initialisieren des Kausalmodells (DE).
        初始化因果模型 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - model_type (str): The type of causal model ('meta' or 'tree') (EN).
                            Тип модели причинности ('meta' или 'tree') (RU).
                            Art des Kausalmodells ('meta' oder 'tree') (DE).
                            因果模型的类型 ('meta' 或 'tree') (ZH).
        """
        if model_type == 'meta':
            self.model = LRSRegressor()
        elif model_type == 'tree':
            from causalml.inference.tree import CausalTreeRegressor
            self.model = CausalTreeRegressor()

    def fit(self, X, treatment, y):
        """
        Train the causal model (EN).
        Обучение причинно-следственной модели (RU).
        Training des Kausalmodells (DE).
        训练因果模型 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - X (array-like): Feature matrix (EN).
                          Матрица признаков (RU).
                          Merkmalsmatrix (DE).
                          特征矩阵 (ZH).
        - treatment (array-like): Treatment variable (EN).
                                  Переменная воздействия (RU).
                                  Behandlungsvariable (DE).
                                  处理变量 (ZH).
        - y (array-like): Outcome variable (EN).
                          Зависимая переменная (RU).
                          Ergebnisvariable (DE).
                          结果变量 (ZH).
        """
        self.model.fit(X, treatment, y)

    def predict(self, X):
        """
        Predict using the causal model (EN).
        Предсказание с использованием причинно-следственной модели (RU).
        Vorhersage mit dem Kausalmodell (DE).
        使用因果模型进行预测 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - X (array-like): Feature matrix for prediction (EN).
                          Матрица признаков для предсказания (RU).
                          Merkmalsmatrix für Vorhersagen (DE).
                          用于预测的特征矩阵 (ZH).

        Returns (EN):
        Возвращает (RU):
        Rückgabe (DE):
        返回 (ZH):
        - predictions (array-like): Predicted causal effects (EN).
                                    Предсказанные причинно-следственные эффекты (RU).
                                    Vorhergesagte Kausalwirkungen (DE).
                                    预测的因果效应 (ZH).
        """
        return self.model.predict(X)

    def feature_importance(self, output_dir):
        """
        Visualize feature importance based on causal analysis (EN).
        Визуализация важности признаков на основе причинно-следственного анализа (RU).
        Visualisierung der Merkmalwichtigkeit auf der Grundlage der Kausalanalyse (DE).
        基于因果分析的特征重要性可视化 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - output_dir (str): Directory to save the visualization (EN).
                            Директория для сохранения визуализации (RU).
                            Verzeichnis zum Speichern der Visualisierung (DE).
                            用于保存可视化的目录 (ZH).

        Returns (EN):
        Возвращает (RU):
        Rückgabe (DE):
        返回 (ZH):
        - output_path (str): Path to the saved feature importance plot (EN).
                             Путь к сохраненному графику важности признаков (RU).
                             Pfad zum gespeicherten Merkmalwichtigkeitsdiagramm (DE).
                             保存的特征重要性图路径 (ZH).
        """
        params = None 
        try:
            if 1 in self.model.models:
                internal_model = self.model.models[1] 
                if hasattr(internal_model, 'coefficients'):
                    params = internal_model.coefficients 
                    print("Model coefficients:", params)  # Перевод: Коэффициенты модели
                else:
                    print("The model does not contain an attribute 'coefficients'.")  # Перевод: Модель не содержит атрибута 'coefficients'.
                    return None
            else:
                print("Model for key 1 not found.")  # Перевод: Модель для ключа 1 не найдена.
                return None
        except AttributeError as e:
            print(f"Failed to extract coefficients. Error: {e}")  # Перевод: Не удалось извлечь коэффициенты. Ошибка:
            return None
        except Exception as e:
            print(f"An error occurred: {str(e)}")  # Перевод: Произошла ошибка:
            return None

        if params is not None:
            plt.figure(figsize=(10, 6))
            sns.barplot(x=np.arange(len(params)), y=params)
            plt.xlabel("Features")  # Перевод: Признаки
            plt.ylabel("Importance (coefficients)")  # Перевод: Важность (коэффициенты)
            plt.title("Feature Importance for Causal Model")  # Перевод: Важность признаков для причинно-следственной модели

            if not os.path.exists(output_dir):
                os.makedirs(output_dir)

            output_path = os.path.join(output_dir, "causal_feature_importance.png")
            plt.savefig(output_path)
            plt.close()
            print(f"Feature importance visualization saved to: {output_path}")  # Перевод: Визуализация важности признаков сохранена в:
            return output_path
        else:
            print("Failed to retrieve feature importance.")  # Перевод: Не удалось получить важность признаков.
            return None


    def plot_causal_effects(self, X, treatment, y, output_dir):
        """
        Visualize causal effects (EN).
        Визуализация причинно-следственных эффектов (RU).
        Visualisierung der Kausalwirkungen (DE).
        因果效应的可视化 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - X (array-like): Feature matrix (EN).
                          Матрица признаков (RU).
                          Merkmalsmatrix (DE).
                          特征矩阵 (ZH).
        - treatment (array-like): Treatment variable (EN).
                                  Переменная воздействия (RU).
                                  Behandlungsvariable (DE).
                                  处理变量 (ZH).
        - y (array-like): Outcome variable (EN).
                          Зависимая переменная (RU).
                          Ergebnisvariable (DE).
                          结果变量 (ZH).
        - output_dir (str): Directory to save the visualization (EN).
                            Директория для сохранения визуализации (RU).
                            Verzeichnis zum Speichern der Visualisierung (DE).
                            用于保存可视化的目录 (ZH).

        Returns (EN):
        Возвращает (RU):
        Rückgabe (DE):
        返回 (ZH):
        - output_path (str): Path to the saved causal effects plot (EN).
                             Путь к сохраненному графику причинно-следственных эффектов (RU).
                             Pfad zum gespeicherten Kausalwirkungsdiagramm (DE).
                             保存的因果效应图路径 (ZH).
        """
        uplift = self.model.predict(X)
        plt.figure(figsize=(10, 6))
        plt.scatter(treatment, uplift, c=y, cmap='bwr', edgecolors='k')
        plt.xlabel("Treatment (Intervention)")  # Перевод: Лечение (Воздействие)
        plt.ylabel("Causal Effect")  # Перевод: Причинно-следственный эффект
        plt.title("Causal Effect Based on the Model")  # Перевод: Причинно-следственный эффект на основе модели
        
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        output_path = os.path.join(output_dir, "causal_effects.png")
        print(f"Causal effects visualization saved to: {output_path}")  # Перевод: Визуализация причинно-следственных эффектов сохранена в:
        plt.savefig(output_path)
        plt.close()
        return output_path
    

    def plot_shap_values(self, X, output_dir):
        """
        Visualize SHAP values to explain the model (EN).
        Визуализация SHAP-значений для объяснения модели (RU).
        Visualisierung von SHAP-Werten zur Erklärung des Modells (DE).
        可视化 SHAP 值以解释模型 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - X (array-like): Feature matrix (EN).
                          Матрица признаков (RU).
                          Merkmalsmatrix (DE).
                          特征矩阵 (ZH).
        - output_dir (str): Directory to save the SHAP plot (EN).
                            Директория для сохранения SHAP-графика (RU).
                            Verzeichnis zum Speichern des SHAP-Diagramms (DE).
                            保存 SHAP 图的目录 (ZH).

        Returns (EN):
        Возвращает (RU):
        Rückgabe (DE):
        返回 (ZH):
        - output_path (str): Path to the saved SHAP plot (EN).
                             Путь к сохраненному SHAP-графику (RU).
                             Pfad zum gespeicherten SHAP-Diagramm (DE).
                             保存的 SHAP 图路径 (ZH).
        """
        explainer = shap.Explainer(self.model.predict, X)
        shap_values = explainer(X)

        plt.figure(figsize=(10, 6))
        shap.summary_plot(shap_values, X, show=False)  # Disable plot display (EN), Отключаем вывод графика (RU), Anzeige des Plots deaktivieren (DE), 禁用图形显示 (ZH)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        output_path = os.path.join(output_dir, "shap_summary_plot.png")
        plt.savefig(output_path, bbox_inches='tight')
        plt.close()
        print(f"SHAP values visualization saved to: {output_path}")  # Перевод: Визуализация SHAP-значений сохранена в:
        return output_path

    def counterfactual_analysis(self, X, treatment, feature_index, new_value):
        """
        Counterfactual analysis: what would happen if one feature was changed (EN).
        Контрфактический анализ: что бы произошло, если бы один из признаков изменился (RU).
        Kontrafaktische Analyse: Was wäre passiert, wenn ein Merkmal geändert worden wäre (DE).
        反事实分析：如果更改了某个特征会发生什么 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - X (array-like): Feature matrix (EN).
                          Матрица признаков (RU).
                          Merkmalsmatrix (DE).
                          特征矩阵 (ZH).
        - treatment (array-like): Treatment variable (EN).
                                  Переменная воздействия (RU).
                                  Behandlungsvariable (DE).
                                  处理变量 (ZH).
        - feature_index (int): Index of the feature to modify (EN).
                               Индекс признака для изменения (RU).
                               Index des zu ändernden Merkmals (DE).
                               要修改的特征的索引 (ZH).
        - new_value (float): New value for the selected feature (EN).
                             Новое значение для выбранного признака (RU).
                             Neuer Wert für das ausgewählte Merkmal (DE).
                             选定特征的新值 (ZH).

        Returns (EN):
        Возвращает (RU):
        Rückgabe (DE):
        返回 (ZH):
        - delta_uplift (array-like): Difference in predicted causal effects (EN).
                                     Разница в предсказанных причинно-следственных эффектах (RU).
                                     Unterschied in den vorhergesagten Kausalwirkungen (DE).
                                     预测的因果效应差异 (ZH).
        """
        X_new = X.copy()
        X_new[:, feature_index] = new_value  # Change feature value (EN), Изменяем значение признака (RU), Merkmalwert ändern (DE), 更改特征值 (ZH)
        
        uplift_original = self.model.predict(X)
        uplift_new = self.model.predict(X_new)
        
        delta_uplift = uplift_new - uplift_original
        return delta_uplift

    def estimate_confidence_intervals(self, X, treatment, n_bootstrap=1000, alpha=0.05):
        """
        Estimate confidence intervals for predictions using bootstrap (EN).
        Оценка доверительных интервалов для предсказаний с использованием бутстрапа (RU).
        Schätzung von Konfidenzintervallen für Vorhersagen mit Bootstrapping (DE).
        使用自举法估计预测的置信区间 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - X (array-like): Feature matrix (EN).
                          Матрица признаков (RU).
                          Merkmalsmatrix (DE).
                          特征矩阵 (ZH).
        - treatment (array-like): Treatment variable (EN).
                                  Переменная воздействия (RU).
                                  Behandlungsvariable (DE).
                                  处理变量 (ZH).
        - n_bootstrap (int): Number of bootstrap samples (EN).
                             Количество бутстрап-выборок (RU).
                             Anzahl der Bootstrap-Stichproben (DE).
                             自举样本的数量 (ZH).
        - alpha (float): Significance level for the confidence interval (EN).
                         Уровень значимости для доверительного интервала (RU).
                         Signifikanzniveau für das Konfidenzintervall (DE).
                         置信区间的显著性水平 (ZH).

        Returns (EN):
        Возвращает (RU):
        Rückgabe (DE):
        返回 (ZH):
        - lower_bound (array-like): Lower bound of the confidence interval (EN).
                                    Нижняя граница доверительного интервала (RU).
                                    Untere Grenze des Konfidenzintervalls (DE).
                                    置信区间的下限 (ZH).
        - upper_bound (array-like): Upper bound of the confidence interval (EN).
                                    Верхняя граница доверительного интервала (RU).
                                    Obere Grenze des Konfidenzintervalls (DE).
                                    置信区间的上限 (ZH).
        """
        predictions = []
        
        for _ in range(n_bootstrap):
            idx = np.random.choice(len(X), len(X), replace=True)
            X_sample = X[idx]
            treatment_sample = treatment[idx]
            pred_sample = self.model.predict(X_sample)
            predictions.append(pred_sample)
        
        predictions = np.array(predictions)
        
        lower_bound = np.percentile(predictions, alpha/2*100, axis=0)
        upper_bound = np.percentile(predictions, (1-alpha/2)*100, axis=0)
        
        return lower_bound, upper_bound

    def analyze_interactions(self, X, treatment, output_dir):
        """
        Analyze interactions between features, selecting the right explainer depending on the model type (EN).
        Анализ взаимодействий между признаками с выбором правильного объяснителя в зависимости от типа модели (RU).
        Analyse der Wechselwirkungen zwischen Merkmalen mit der Auswahl des richtigen Erklärers je nach Modelltyp (DE).
        分析特征之间的相互作用，根据模型类型选择正确的解释器 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - X (array-like): Feature matrix (EN).
                          Матрица признаков (RU).
                          Merkmalsmatrix (DE).
                          特征矩阵 (ZH).
        - treatment (array-like): Treatment variable (EN).
                                  Переменная воздействия (RU).
                                  Behandlungsvariable (DE).
                                  处理变量 (ZH).
        - output_dir (str): Directory to save the visualization (EN).
                            Директория для сохранения визуализации (RU).
                            Verzeichnis zum Speichern der Visualisierung (DE).
                            用于保存可视化的目录 (ZH).

        Returns (EN):
        Возвращает (RU):
        Rückgabe (DE):
        返回 (ZH):
        - output_path (str): Path to the saved interaction plot (EN).
                             Путь к сохраненному графику взаимодействий (RU).
                             Pfad zum gespeicherten Interaktionsdiagramm (DE).
                             保存的相互作用图路径 (ZH).
        """
        model_type = type(self.model)

        if model_type.__name__ == 'CausalTreeRegressor':
            explainer = shap.TreeExplainer(self.model)
            shap_interaction_values = explainer.shap_interaction_values(X)
        elif model_type.__name__ == 'LRSRegressor':
            explainer = shap.Explainer(self.model.predict, X)
            shap_interaction_values = explainer(X)
        else:
            print(f"Model type '{model_type.__name__}' is not supported for interaction analysis.")  # Перевод: Тип модели '{model_type.__name__}' не поддерживается для анализа взаимодействий.
            return None

        plt.figure(figsize=(10, 6))
        shap.summary_plot(shap_interaction_values, X, show=False)
        plt.title("Feature Interactions")  # Перевод: Взаимодействие признаков

        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        output_path = os.path.join(output_dir, "interaction_plot.png")
        plt.savefig(output_path, bbox_inches='tight')
        plt.close()
        print(f"Feature interactions visualization saved to: {output_path}")  # Перевод: Взаимодействие признаков сохранено в:
        return output_path
    
    def generate_report(self, X_test, treatment_test, y_test, output_dir):
        """
        Generate an Excel and Word report containing all key results and visualizations (EN).
        Генерация отчета в формате Excel и Word, содержащего все важные результаты и визуализации (RU).
        Generieren eines Berichts im Excel- und Word-Format, der alle wichtigen Ergebnisse und Visualisierungen enthält (DE).
        生成包含所有关键结果和可视化内容的 Excel 和 Word 报告 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - X_test (array-like): Test feature matrix (EN).
                               Тестовая матрица признаков (RU).
                               Test-Merkmalsmatrix (DE).
                               测试特征矩阵 (ZH).
        - treatment_test (array-like): Test treatment variable (EN).
                                       Тестовая переменная воздействия (RU).
                                       Test-Behandlungsvariable (DE).
                                       测试处理变量 (ZH).
        - y_test (array-like): Test outcome variable (EN).
                               Тестовая зависимая переменная (RU).
                               Test-Ergebnisvariable (DE).
                               测试结果变量 (ZH).
        - output_dir (str): Directory to save the report and visualizations (EN).
                            Директория для сохранения отчета и визуализаций (RU).
                            Verzeichnis zum Speichern des Berichts und der Visualisierungen (DE).
                            保存报告和可视化的目录 (ZH).

        Returns (EN):
        Возвращает (RU):
        Rückgabe (DE):
        返回 (ZH):
        - excel_report_path (str): Path to the Excel report (EN).
                                   Путь к отчету в формате Excel (RU).
                                   Pfad zum Excel-Bericht (DE).
                                   Excel 报告的路径 (ZH).
        - word_report_path (str): Path to the Word report (EN).
                                  Путь к отчету в формате Word (RU).
                                  Pfad zum Word-Bericht (DE).
                                  Word 报告的路径 (ZH).
        """
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        feature_importance_path = self.feature_importance(output_dir)
        causal_effects_path = self.plot_causal_effects(X_test, treatment_test, y_test, output_dir)
        shap_path = self.plot_shap_values(X_test, output_dir)
        interaction_path = self.analyze_interactions(X_test, treatment_test, output_dir)

        excel_report_path = os.path.join(output_dir, "causal_analysis_report.xlsx")

        # Coefficients and feature importances
        if hasattr(self.model.models[1], 'coefficients'):
            coefficients = self.model.models[1].coefficients
            df_importance = pd.DataFrame({
                'Feature': [f'Feature_{i}' for i in range(len(coefficients))],
                'Coefficient': coefficients
            })
            with pd.ExcelWriter(excel_report_path, engine='openpyxl') as writer:
                df_importance.to_excel(writer, sheet_name='Feature Importance', index=False)

                # Add the counterfactual analysis result
                delta_uplift = self.counterfactual_analysis(X_test, treatment_test, feature_index=0, new_value=1.0)
                if delta_uplift.ndim > 1:
                    delta_uplift = delta_uplift.flatten()

                df_counterfactual = pd.DataFrame({
                    'Observation': np.arange(len(delta_uplift)),
                    'Delta Uplift': delta_uplift
                })
                df_counterfactual.to_excel(writer, sheet_name='Counterfactual Analysis', index=False)

                # Add confidence intervals
                lower_bound, upper_bound = self.estimate_confidence_intervals(X_test, treatment_test)
                df_confidence_intervals = pd.DataFrame({
                    'Lower Bound': lower_bound.flatten(),
                    'Upper Bound': upper_bound.flatten()
                })
                df_confidence_intervals.to_excel(writer, sheet_name='Confidence Intervals', index=False)

                # Save paths to visualizations in Excel
                df_visualization_paths = pd.DataFrame({
                    'Visualization': ['Feature Importance', 'Causal Effects', 'SHAP Summary', 'Interactions'],
                    'File Path': [feature_importance_path, causal_effects_path, shap_path, interaction_path]
                })
                df_visualization_paths.to_excel(writer, sheet_name='Visualizations', index=False)

        wb = load_workbook(excel_report_path)
        sheet_name = 'Graphs'
        if sheet_name not in wb.sheetnames:
            ws = wb.create_sheet(sheet_name)
        else:
            ws = wb[sheet_name]

        img_feature_importance = Image(feature_importance_path)
        img_causal_effects = Image(causal_effects_path)
        img_shap = Image(shap_path)
        img_interaction = Image(interaction_path)

        ws.add_image(img_feature_importance, 'A1')
        ws.add_image(img_causal_effects, 'A20')
        ws.add_image(img_shap, 'A40')
        ws.add_image(img_interaction, 'A60')

        wb.save(excel_report_path)

        word_report_path = os.path.join(output_dir, "causal_analysis_report.docx")
        doc = Document()

        doc.add_heading('Causal Analysis Report', 0)
        doc.add_heading('Feature Importance', level=1)
        doc.add_paragraph('Feature importance calculated from the causal model:')

        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Feature'
        hdr_cells[1].text = 'Coefficient'
        for i, coef in enumerate(coefficients):
            row_cells = table.add_row().cells
            row_cells[0].text = f'Feature_{i}'
            row_cells[1].text = str(coef)

        doc.add_heading('Visualizations', level=1)
        doc.add_paragraph('Below are visualizations generated by the model:')

        doc.add_paragraph('Feature Importance:')
        doc.add_picture(feature_importance_path, width=Inches(4))

        doc.add_paragraph('Causal Effects:')
        doc.add_picture(causal_effects_path, width=Inches(4))

        doc.add_paragraph('SHAP Summary:')
        doc.add_picture(shap_path, width=Inches(4))

        doc.add_paragraph('Feature Interactions:')
        doc.add_picture(interaction_path, width=Inches(4))

        doc.add_heading('Counterfactual Analysis', level=1)
        doc.add_paragraph('Changes in predictions when the first feature is changed to 1.0:')
        
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Observation'
        hdr_cells[1].text = 'Delta Uplift'
        for i, delta in enumerate(delta_uplift):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = str(delta)

        doc.add_heading('Confidence Intervals', level=1)
        doc.add_paragraph('Confidence intervals calculated for the predictions:')

        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Lower Bound'
        hdr_cells[1].text = 'Upper Bound'
        for lb, ub in zip(lower_bound, upper_bound):
            row_cells = table.add_row().cells
            row_cells[0].text = str(lb)
            row_cells[1].text = str(ub)

        doc.save(word_report_path)
        print(f"Word report saved to: {word_report_path}")  # Перевод: Отчет в формате Word сохранен в:

        return excel_report_path, word_report_path
