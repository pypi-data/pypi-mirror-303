import torch
import torch.nn as nn
import torch.nn.functional as F
import matplotlib.pyplot as plt
from torch.utils.tensorboard import SummaryWriter
from docx import Document
from docx.shared import Inches
import os

# Create directory for logs (EN), Создание директории для логов (RU), Erstellen eines Verzeichnisses für Protokolle (DE), 创建日志目录 (ZH)
os.makedirs("/logs", exist_ok=True)
writer = SummaryWriter(log_dir="/logs")

class AdaptiveActivation(nn.Module):
    def __init__(self, num_neurons, num_functions=5):
        """
        Adaptive activation module allows dynamic selection of activation functions per neuron (EN).
        Модуль адаптивной активации позволяет динамически выбирать функции активации для каждого нейрона (RU).
        Das adaptive Aktivierungsmodul ermöglicht eine dynamische Auswahl von Aktivierungsfunktionen für jedes Neuron (DE).
        自适应激活模块允许为每个神经元动态选择激活函数 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - num_neurons (int): Number of neurons in the layer (EN).
                             Количество нейронов в слое (RU).
                             Anzahl der Neuronen in der Schicht (DE).
                             层中的神经元数量 (ZH).
        - num_functions (int): Number of possible activation functions (default: 5) (EN).
                               Количество возможных функций активации (по умолчанию: 5) (RU).
                               Anzahl der möglichen Aktivierungsfunktionen (Standard: 5) (DE).
                               可能的激活函数数量 (默认：5) (ZH).
        """
        super(AdaptiveActivation, self).__init__()
        # Learnable weights for combining activation functions (EN)
        # Обучаемые веса для комбинирования функций активации (RU)
        # Lernbare Gewichte zur Kombination von Aktivierungsfunktionen (DE)
        # 可学习的权重，用于组合激活函数 (ZH)
        self.weights = nn.Parameter(torch.randn(num_neurons, num_functions, requires_grad=True))
        # Attention weights for controlling the influence of activations (EN)
        # Веса внимания для управления влиянием активаций (RU)
        # Aufmerksamkeitsgewichte zur Steuerung des Einflusses von Aktivierungen (DE)
        # 注意力权重，用于控制激活的影响 (ZH)
        self.attention_weights = nn.Parameter(torch.ones(num_neurons, requires_grad=True)) 
        # List of possible activation functions (EN), Список возможных функций активации (RU), Liste möglicher Aktivierungsfunktionen (DE), 可能的激活函数列表 (ZH)
        self.activations = [F.relu, torch.sigmoid, torch.tanh, F.elu, lambda x: x * torch.sigmoid(x)]

    def forward(self, x):
        """
        Forward pass for applying weighted activation functions (EN).
        Прямой проход для применения взвешенных функций активации (RU).
        Vorwärtsdurchlauf zur Anwendung gewichteter Aktivierungsfunktionen (DE).
        应用加权激活函数的前向传递 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - x (torch.Tensor): Input tensor (EN).
                            Входной тензор (RU).
                            Eingabetensor (DE).
                            输入张量 (ZH).

        Returns (EN):
        Возвращает (RU):
        Rückgabe (DE):
        返回 (ZH):
        - output (torch.Tensor): Output tensor after applying weighted activations (EN).
                                 Выходной тензор после применения взвешенных активаций (RU).
                                 Ausgangstensor nach Anwendung gewichteter Aktivierungen (DE).
                                 应用加权激活后的输出张量 (ZH).
        """
        normalized_weights = F.softmax(self.weights, dim=1)
        attention = F.sigmoid(self.attention_weights).unsqueeze(0).expand(x.size(0), -1)
        
        output = torch.zeros_like(x)
        for i, activation in enumerate(self.activations):
            activation_result = activation(x)
            weight = normalized_weights[:, i].unsqueeze(0).expand_as(activation_result)
            attention_expanded = attention.expand_as(activation_result)
            # Applying weighted activation with attention (EN)
            # Применение взвешенной активации с вниманием (RU)
            # Anwendung gewichteter Aktivierung mit Aufmerksamkeit (DE)
            # 应用带有注意力的加权激活 (ZH)
            output += weight * activation_result * attention_expanded
        
        return output


class DropConnectActivation(AdaptiveActivation):
    def forward(self, x, drop_prob=0.5):
        """
        Forward pass with DropConnect regularization (EN).
        Прямой проход с регуляризацией DropConnect (RU).
        Vorwärtsdurchlauf mit DropConnect-Regularisierung (DE).
        带有 DropConnect 正则化的前向传递 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - x (torch.Tensor): Input tensor (EN).
                            Входной тензор (RU).
                            Eingabetensor (DE).
                            输入张量 (ZH).
        - drop_prob (float): DropConnect probability (default: 0.5) (EN).
                             Вероятность DropConnect (по умолчанию: 0.5) (RU).
                             DropConnect-Wahrscheinlichkeit (Standard: 0.5) (DE).
                             DropConnect 概率 (默认: 0.5) (ZH).

        Returns (EN):
        Возвращает (RU):
        Rückgabe (DE):
        返回 (ZH):
        - output (torch.Tensor): Output tensor after applying DropConnect (EN).
                                 Выходной тензор после применения DropConnect (RU).
                                 Ausgangstensor nach Anwendung von DropConnect (DE).
                                 应用 DropConnect 后的输出张量 (ZH).
        """
        normalized_weights = F.softmax(self.weights, dim=1)
        output = torch.zeros_like(x)
        for i, activation in enumerate(self.activations):
            # DropConnect mask for randomly zeroing out connections (EN)
            # Маска DropConnect для случайного обнуления связей (RU)
            # DropConnect-Maske zum zufälligen Nullsetzen von Verbindungen (DE)
            # DropConnect 掩码，用于随机清零连接 (ZH)
            mask = (torch.rand_like(x) > drop_prob).float()
            output += normalized_weights[:, i].unsqueeze(1) * activation(x) * mask
        return output


class AdaptiveLayer(nn.Module):
    def __init__(self, input_size, output_size, drop_connect=False):
        """
        Adaptive layer with optional DropConnect (EN).
        Адаптивный слой с опциональным DropConnect (RU).
        Adaptiver Layer mit optionalem DropConnect (DE).
        带有可选 DropConnect 的自适应层 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - input_size (int): Size of the input features (EN).
                            Размер входных признаков (RU).
                            Größe der Eingabemerkmale (DE).
                            输入特征的大小 (ZH).
        - output_size (int): Size of the output features (EN).
                             Размер выходных признаков (RU).
                             Größe der Ausgangsmerkmale (DE).
                             输出特征的大小 (ZH).
        - drop_connect (bool): Whether to use DropConnect regularization (default: False) (EN).
                               Использовать ли регуляризацию DropConnect (по умолчанию: False) (RU).
                               Ob DropConnect-Regularisierung verwendet werden soll (Standard: False) (DE).
                               是否使用 DropConnect 正则化 (默认: False) (ZH).
        """
        super(AdaptiveLayer, self).__init__()
        self.fc = nn.Linear(input_size, output_size)
        # If DropConnect is enabled, use DropConnectActivation (EN)
        # Если DropConnect включен, используйте DropConnectActivation (RU)
        # Wenn DropConnect aktiviert ist, verwenden Sie DropConnectActivation (DE)
        # 如果启用 DropConnect，请使用 DropConnectActivation (ZH)
        if drop_connect:
            self.adaptive_activation = DropConnectActivation(output_size)
        else:
            self.adaptive_activation = AdaptiveActivation(output_size)

    def forward(self, x):
        """
        Forward pass through the adaptive layer (EN).
        Прямой проход через адаптивный слой (RU).
        Vorwärtsdurchlauf durch den adaptiven Layer (DE).
        通过自适应层的前向传递 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - x (torch.Tensor): Input tensor (EN).
                            Входной тензор (RU).
                            Eingabetensor (DE).
                            输入张量 (ZH).

        Returns (EN):
        Возвращает (RU):
        Rückgabe (DE):
        返回 (ZH):
        - output (torch.Tensor): Output tensor after applying the adaptive activation (EN).
                                 Выходной тензор после применения адаптивной активации (RU).
                                 Ausgangstensor nach Anwendung der adaptiven Aktivierung (DE).
                                 应用自适应激活后的输出张量 (ZH).
        """
        x = self.fc(x)
        x = self.adaptive_activation(x)
        return x


class MultiTaskLearningNet(nn.Module):
    def __init__(self, input_size, hidden_size, output_size_regression, output_size_classification):
        """
        Multi-task learning network for regression and classification (EN).
        Модель мультитаскинга для регрессии и классификации (RU).
        Mehrzwecknetzwerk für Regression und Klassifikation (DE).
        用于回归和分类的多任务学习网络 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - input_size (int): Input feature size (EN).
                            Размер входных признаков (RU).
                            Größe der Eingabemerkmale (DE).
                            输入特征的大小 (ZH).
        - hidden_size (int): Hidden layer size (EN).
                             Размер скрытого слоя (RU).
                             Größe der verborgenen Schicht (DE).
                             隐藏层的大小 (ZH).
        - output_size_regression (int): Size of the regression output (EN).
                                        Размер выходных данных для регрессии (RU).
                                        Größe der Regressionsausgabe (DE).
                                        回归输出的大小 (ZH).
        - output_size_classification (int): Size of the classification output (EN).
                                            Размер выходных данных для классификации (RU).
                                            Größe der Klassifikationsausgabe (DE).
                                            分类输出的大小 (ZH).
        """
        super(MultiTaskLearningNet, self).__init__()
        self.layer1 = AdaptiveLayer(input_size, hidden_size)
        self.regression_head = nn.Linear(hidden_size, output_size_regression)
        self.classification_head = nn.Linear(hidden_size, output_size_classification)
        self.loss_history = [] 

    def forward(self, x):
        """
        Forward pass for both regression and classification tasks (EN).
        Прямой проход для задач регрессии и классификации (RU).
        Vorwärtsdurchlauf für Regressions- und Klassifikationsaufgaben (DE).
        回归和分类任务的前向传递 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - x (torch.Tensor): Input tensor (EN).
                            Входной тензор (RU).
                            Eingabetensor (DE).
                            输入张量 (ZH).

        Returns (EN):
        Возвращает (RU):
        Rückgabe (DE):
        返回 (ZH):
        - regression_output (torch.Tensor): Output for the regression task (EN).
                                            Выходные данные для задачи регрессии (RU).
                                            Ausgangsdaten für die Regressionsaufgabe (DE).
                                            回归任务的输出 (ZH).
        - classification_output (torch.Tensor): Output for the classification task (EN).
                                                Выходные данные для задачи классификации (RU).
                                                Ausgangsdaten für die Klassifikationsaufgabe (DE).
                                                分类任务的输出 (ZH).
        """
        x = self.layer1(x)
        regression_output = self.regression_head(x)
        classification_output = self.classification_head(x)
        return regression_output, classification_output

    def log_metrics(self, loss_regression, loss_classification, epoch):
        """
        Log regression and classification losses to TensorBoard (EN).
        Логирование потерь регрессии и классификации в TensorBoard (RU).
        Protokollieren von Regressions- und Klassifikationsverlusten in TensorBoard (DE).
        将回归和分类损失记录到 TensorBoard (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - loss_regression (float): Regression loss (EN).
                                   Потери регрессии (RU).
                                   Regressionsverlust (DE).
                                   回归损失 (ZH).
        - loss_classification (float): Classification loss (EN).
                                       Потери классификации (RU).
                                       Klassifikationsverlust (DE).
                                       分类损失 (ZH).
        - epoch (int): Current epoch number (EN).
                       Номер текущей эпохи (RU).
                       Aktuelle Epochenzahl (DE).
                       当前周期数 (ZH).
        """
        total_loss = loss_regression + loss_classification
        self.loss_history.append(total_loss.item()) 
        writer.add_scalar('Loss/Regression', loss_regression, epoch)
        writer.add_scalar('Loss/Classification', loss_classification, epoch)

    def log_activation_weights(self, epoch):
        """
        Log activation weights to TensorBoard (EN).
        Логирование весов активации в TensorBoard (RU).
        Protokollieren von Aktivierungsgewichten in TensorBoard (DE).
        将激活权重记录到 Tensor Board (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - epoch (int): Current epoch number (EN).
                       Номер текущей эпохи (RU).
                       Aktuelle Epochenzahl (DE).
                       当前周期数 (ZH).
        """
        for name, param in self.named_parameters():
            if "weights" in name:
                writer.add_histogram(f'Weights/{name}', param, epoch)

    def regularization_loss(self):
        """
        Compute L1 regularization loss (EN).
        Вычисление потерь L1-регуляризации (RU).
        Berechnung des L1-Regularisierungsverlusts (DE).
        计算 L1 正则化损失 (ZH).

        Returns (EN):
        Возвращает (RU):
        Rückgabe (DE):
        返回 (ZH):
        - reg_loss (float): L1 regularization loss (EN).
                            Потери L1-регуляризации (RU).
                            L1-Regularisierungsverlust (DE).
                            L1 正则化损失 (ZH).
        """
        reg_loss = 0.0
        for param in self.parameters():
            if param.requires_grad:
                reg_loss += torch.sum(torch.abs(param)) 
        return reg_loss


def plot_loss_history(loss_history):
    """
    Plot the loss history (EN).
    Построение графика истории потерь (RU).
    Verlauf der Verlusthistorie anzeigen (DE).
    绘制损失历史 (ZH).

    Parameters (EN):
    Параметры (RU):
    Parameter (DE):
    参数 (ZH):
    - loss_history (list of float): List of loss values across epochs (EN).
                                    Список значений потерь по эпохам (RU).
                                    Liste der Verlustwerte über Epochen (DE).
                                    各周期的损失值列表 (ZH).
    """
    plt.plot(range(len(loss_history)), loss_history, label='Loss')
    plt.title('Loss History')
    plt.xlabel('Epoch')
    plt.ylabel('Loss')
    plt.legend()
    plt.show()

def save_loss_plot(loss_history, filename="loss_plot.png"):
    """
    Save the loss history plot as an image file (EN).
    Сохранить график истории потерь как файл изображения (RU).
    Speichern Sie den Verlustverlauf als Bilddatei (DE).
    将损失历史图保存为图像文件 (ZH).

    Parameters (EN):
    Параметры (RU):
    Parameter (DE):
    参数 (ZH):
    - loss_history (list of float): List of loss values across epochs (EN).
                                    Список значений потерь по эпохам (RU).
                                    Liste der Verlustwerte über Epochen (DE).
                                    各周期的损失值列表 (ZH).
    - filename (str): Path to save the image file (EN).
                      Путь для сохранения файла изображения (RU).
                      Pfad zum Speichern der Bilddatei (DE).
                      保存图像文件的路径 (ZH).
    """
    plt.figure()
    plt.plot(range(len(loss_history)), loss_history, label='Loss')
    plt.title('Loss History')
    plt.xlabel('Epoch')
    plt.ylabel('Loss')
    plt.legend()
    plt.savefig(filename) 
    plt.close()

def generate_report(model, loss_history, report_filename="training_report.docx"):
    """
    Generate a training report in Word format, including loss history and model details (EN).
    Генерация отчета о тренировке в формате Word, включающего историю потерь и информацию о модели (RU).
    Erstellen Sie einen Trainingsbericht im Word-Format, der den Verlustverlauf und Modelldetails enthält (DE).
    生成包含损失历史和模型详细信息的 Word 格式训练报告 (ZH).

    Parameters (EN):
    Параметры (RU):
    Parameter (DE):
    参数 (ZH):
    - model (nn.Module): PyTorch model (EN).
                         PyTorch модель (RU).
                         PyTorch-Modell (DE).
                         PyTorch 模型 (ZH).
    - loss_history (list of float): List of loss values across epochs (EN).
                                    Список значений потерь по эпохам (RU).
                                    Liste der Verlustwerte über Epochen (DE).
                                    各周期的损失值列表 (ZH).
    - report_filename (str): Path to save the report (EN).
                             Путь для сохранения отчета (RU).
                             Pfad zum Speichern des Berichts (DE).
                             保存报告的路径 (ZH).
    """
    doc = Document()
    doc.add_heading('Training Report', 0)
    
    doc.add_heading('Model Description', level=1)
    model_description = f"The model has {sum(p.numel() for p in model.parameters() if p.requires_grad)} trainable parameters."
    doc.add_paragraph(model_description)
    
    doc.add_heading('Training Metrics', level=1)
    final_loss = loss_history[-1]
    doc.add_paragraph(f"Final Loss: {final_loss}")
    
    plot_filename = "loss_plot.png"
    save_loss_plot(loss_history, filename=plot_filename)
    
    doc.add_heading('Loss History', level=1)
    doc.add_paragraph('The following plot shows the loss history over epochs:')
    doc.add_picture(plot_filename, width=Inches(5.0))
    
    report_path = os.path.join(os.getcwd(), report_filename)
    doc.save(report_path)
    print(f"Report saved to {report_path}")
