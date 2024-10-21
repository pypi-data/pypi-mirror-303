from .model_tuner import ModelTuner
from .causal_model import IljicevsCausalModel
from sklearn.model_selection import cross_val_score
from sklearn.metrics import accuracy_score, classification_report
import os
import pandas as pd
from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook
from openpyxl.drawing.image import Image as ExcelImage

class IljicevsAnsambleModel:
    def __init__(self, models=None, param_grids=None, search_method="grid", use_causal_model=False):
        """
        Initialization of the ensemble model with options for hyperparameter search method and causal model (EN).
        Инициализация ансамбля моделей с возможностью выбора метода поиска гиперпараметров и причинно-следственной модели (RU).
        Initialisierung des Ensemble-Modells mit Optionen für die Methode zur Suche nach Hyperparametern und Kausalmodell (DE).
        使用超参数搜索方法和因果模型选项的集成模型初始化 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - models (list): List of machine learning models to be included in the ensemble (EN).
                         Список моделей машинного обучения для включения в ансамбль (RU).
                         Liste von Modellen des maschinellen Lernens, die im Ensemble enthalten sind (DE).
                         要包含在集成中的机器学习模型列表 (ZH).
        - param_grids (list of dicts): List of hyperparameter grids for tuning each model (EN).
                                       Список сеток гиперпараметров для настройки каждой модели (RU).
                                       Liste von Hyperparameter-Rastern zur Abstimmung jedes Modells (DE).
                                       调整每个模型的超参数网格列表 (ZH).
        - search_method (str): Search method for hyperparameter tuning ('grid', 'random', etc.) (EN).
                               Метод поиска гиперпараметров ('grid', 'random' и т.д.) (RU).
                               Suchmethode für die Hyperparameter-Abstimmung ('grid', 'random' usw.) (DE).
                               超参数调整的搜索方法（'grid'，'random' 等）(ZH).
        - use_causal_model (bool): Whether to use the causal model for training (EN).
                                   Использовать ли причинно-следственную модель для обучения (RU).
                                   Ob das Kausalmodell für das Training verwendet werden soll (DE).
                                   是否使用因果模型进行训练 (ZH).
        """
        self.models = models
        self.param_grids = param_grids
        self.search_method = search_method
        self.best_models = None
        self.selected_models = None
        self.use_causal_model = use_causal_model
        if use_causal_model:
            self.causal_model = IljicevsCausalModel()

    def fit(self, X_train, y_train, treatment=None):
        """
        Train the selected models on the training dataset (EN).
        Обучение выбранных моделей на обучающем наборе данных (RU).
        Training der ausgewählten Modelle auf dem Trainingsdatensatz (DE).
        在训练数据集上训练选定的模型 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - X_train (array-like): Feature matrix for training (EN).
                                Матрица признаков для обучения (RU).
                                Merkmalsmatrix zum Training (DE).
                                用于训练的特征矩阵 (ZH).
        - y_train (array-like): Target variable for training (EN).
                                Целевая переменная для обучения (RU).
                                Zielvariable zum Training (DE).
                                用于训练的目标变量 (ZH).
        - treatment (array-like, optional): Treatment variable for causal model (optional) (EN).
                                            Переменная воздействия для причинно-следственной модели (необязательно) (RU).
                                            Behandlungsvariable für das Kausalmodell (optional) (DE).
                                            因果模型的处理变量（可选） (ZH).
        """
        if self.use_causal_model and treatment is not None:
            self.causal_model.fit(X_train, treatment, y_train)
        for model in self.selected_models:
            model.fit(X_train, y_train)

    def score(self, X_test, y_test):
        """
        Evaluate the accuracy of the ensemble model on the test dataset (EN).
        Оценка точности ансамбля моделей на тестовых данных (RU).
        Bewertung der Genauigkeit des Ensemble-Modells auf dem Testdatensatz (DE).
        在测试数据集上评估集成模型的准确性 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - X_test (array-like): Feature matrix for testing (EN).
                               Матрица признаков для тестирования (RU).
                               Merkmalsmatrix zum Testen (DE).
                               用于测试的特征矩阵 (ZH).
        - y_test (array-like): Target variable for testing (EN).
                               Целевая переменная для тестирования (RU).
                               Zielvariable zum Testen (DE).
                               用于测试的目标变量 (ZH).

        Returns (EN):
        Возвращает (RU):
        Rückgabe (DE):
        返回 (ZH):
        - accuracy (float): Accuracy score of the ensemble model on the test dataset (EN).
                            Оценка точности ансамбля моделей на тестовом наборе данных (RU).
                            Genauigkeitswert des Ensemble-Modells auf dem Testdatensatz (DE).
                            集成模型在测试数据集上的准确性评分 (ZH).
        """
        predictions = [model.predict(X_test) for model in self.selected_models]
        return accuracy_score(y_test, predictions)

    def generate_report(self, X_test, y_test, output_dir):
        """
        Generate a classification report and save it to Word and Excel (EN).
        Генерация отчёта с результатами классификации и сохранение в Word и Excel (RU).
        Generieren eines Klassifikationsberichts und Speichern in Word und Excel (DE).
        生成分类报告并保存为 Word 和 Excel 文件 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - X_test (array-like): Feature matrix for testing (EN).
                               Матрица признаков для тестирования (RU).
                               Merkmalsmatrix zum Testen (DE).
                               用于测试的特征矩阵 (ZH).
        - y_test (array-like): Target variable for testing (EN).
                               Целевая переменная для тестирования (RU).
                               Zielvariable zum Testen (DE).
                               用于测试的目标变量 (ZH).
        - output_dir (str): Directory to save the generated reports (EN).
                            Директория для сохранения сгенерированных отчётов (RU).
                            Verzeichnis zum Speichern der generierten Berichte (DE).
                            保存生成报告的目录 (ZH).

        Returns (EN):
        Возвращает (RU):
        Rückgabe (DE):
        返回 (ZH):
        - word_path (str): Path to the saved Word report (EN).
                           Путь к сохранённому отчёту в формате Word (RU).
                           Pfad zum gespeicherten Word-Bericht (DE).
                           保存的 Word 报告的路径 (ZH).
        - excel_path (str): Path to the saved Excel report (EN).
                            Путь к сохранённому отчёту в формате Excel (RU).
                            Pfad zum gespeicherten Excel-Bericht (DE).
                            保存的 Excel 报告的路径 (ZH).
        """
        report = classification_report(y_test, [model.predict(X_test) for model in self.selected_models], output_dict=True)
        df = pd.DataFrame(report).transpose()

        # Сохранение отчёта в Excel (RU), Save report in Excel (EN), Speichern des Berichts in Excel (DE), 将报告保存为 Excel (ZH)
        excel_path = os.path.join(output_dir, "classification_report.xlsx")
        df.to_excel(excel_path)

        # Сохранение отчёта в Word (RU), Save report in Word (EN), Speichern des Berichts in Word (DE), 将报告保存为 Word (ZH)
        doc = Document()
        doc.add_heading('Classification Report', 0)
        doc.add_paragraph(str(df))
        word_path = os.path.join(output_dir, "classification_report.docx")
        doc.save(word_path)

        return word_path, excel_path

    def fit_and_report(self, X_train, y_train, X_test, y_test, output_dir):
        """
        Full cycle: model training, prediction, and report generation (EN).
        Полный цикл: обучение моделей, предсказание и генерация отчётов (RU).
        Vollständiger Zyklus: Modelltraining, Vorhersage und Berichterstellung (DE).
        完整流程：模型训练、预测和报告生成 (ZH).

        Parameters (EN):
        Параметры (RU):
        Parameter (DE):
        参数 (ZH):
        - X_train (array-like): Feature matrix for training (EN).
                                Матрица признаков для обучения (RU).
                                Merkmalsmatrix zum Training (DE).
                                用于训练的特征矩阵 (ZH).
        - y_train (array-like): Target variable for training (EN).
                                Целевая переменная для обучения (RU).
                                Zielvariable zum Training (DE).
                                用于训练的目标变量 (ZH).
        - X_test (array-like): Feature matrix for testing (EN).
                               Матрица признаков для тестирования (RU).
                               Merkmalsmatrix zum Testen (DE).
                               用于测试的特征矩阵 (ZH).
        - y_test (array-like): Target variable for testing (EN).
                               Целевая переменная для тестирования (RU).
                               Zielvariable zum Testen (DE).
                               用于测试的目标变量 (ZH).
        - output_dir (str): Directory to save the generated reports (EN).
                            Директория для сохранения сгенерированных отчётов (RU).
                            Verzeichnis zum Speichern der generierten Berichte (DE).
                            保存生成报告的目录 (ZH).

        Returns (EN):
        Возвращает (RU):
        Rückgabe (DE):
        返回 (ZH):
        - word_path (str): Path to the saved Word report (EN).
                           Путь к сохранённому отчёту в формате Word (RU).
                           Pfad zum gespeicherten Word-Bericht (DE).
                           保存的 Word 报告的路径 (ZH).
        - excel_path (str): Path to the saved Excel report (EN).
                            Путь к сохранённому отчёту в формате Excel (RU).
                            Pfad zum gespeicherten Excel-Bericht (DE).
                            保存的 Excel 报告的路径 (ZH).
        """
        self.fit(X_train, y_train)
        return self.generate_report(X_test, y_test, output_dir)
