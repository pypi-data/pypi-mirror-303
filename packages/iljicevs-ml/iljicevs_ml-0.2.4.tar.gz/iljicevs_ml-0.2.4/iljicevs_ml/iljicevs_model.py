from .model_tuner import ModelTuner
from .causal_model import IljicevsCausalModel
from sklearn.model_selection import cross_val_score
from sklearn.metrics import accuracy_score, classification_report
import os
import pandas as pd
from docx import Document
from docx.shared import Inches
from openpyxl import load_workbook
from openpyxl.drawing.image import Image as ExcelImage

class IljicevsAnsambleModel:
    def __init__(self, models=None, param_grids=None, search_method="grid", use_causal_model=False):
        """
        Инициализация ансамбля моделей с возможностью выбора метода поиска гиперпараметров и причинно-следственной модели.
        """
        self.models = models
        self.param_grids = param_grids
        self.search_method = search_method
        self.best_models = None
        self.selected_models = None
        self.use_causal_model = use_causal_model
        if use_causal_model:
            self.causal_model = IljicevsCausalModel()

    def fit(self, X_train, y_train, treatment=None):
        """
        Обучение выбранных моделей на обучающем наборе данных.
        """
        if self.use_causal_model and treatment is not None:
            self.causal_model.fit(X_train, treatment, y_train)
        for model in self.selected_models:
            model.fit(X_train, y_train)

    def score(self, X_test, y_test):
        """
        Оценка точности ансамбля моделей на тестовых данных.
        """
        predictions = [model.predict(X_test) for model in self.selected_models]
        return accuracy_score(y_test, predictions)

    def generate_report(self, X_test, y_test, output_dir):
        """
        Генерация отчёта с результатами классификации и сохранение в Word и Excel.
        """
        report = classification_report(y_test, [model.predict(X_test) for model in self.selected_models], output_dict=True)
        df = pd.DataFrame(report).transpose()

        # Сохранение отчёта в Excel
        excel_path = os.path.join(output_dir, "classification_report.xlsx")
        df.to_excel(excel_path)

        # Сохранение отчёта в Word
        doc = Document()
        doc.add_heading('Classification Report', 0)
        doc.add_paragraph(str(df))
        word_path = os.path.join(output_dir, "classification_report.docx")
        doc.save(word_path)

        return word_path, excel_path

    def fit_and_report(self, X_train, y_train, X_test, y_test, output_dir):
        """
        Полный цикл: обучение моделей, предсказание и генерация отчётов.
        """
        self.fit(X_train, y_train)
        return self.generate_report(X_test, y_test, output_dir)
