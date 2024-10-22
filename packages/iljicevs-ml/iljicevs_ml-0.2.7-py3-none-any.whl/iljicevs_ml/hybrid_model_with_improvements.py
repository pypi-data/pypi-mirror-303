import numpy as np
import matplotlib.pyplot as plt
import os
from docx import Document
from openpyxl import Workbook

# LSTMCell
# English: This class implements an LSTM cell for sequence processing.
# Russian: Этот класс реализует ячейку LSTM для обработки последовательностей.
# German: Diese Klasse implementiert eine LSTM-Zelle für die Verarbeitung von Sequenzen.
# Chinese: 该类实现了用于序列处理的LSTM单元.
class LSTMCell:
    def __init__(self, input_dim, hidden_dim):
        """
        English: Initialize the LSTM cell with input and hidden dimensions.
        Russian: Инициализация ячейки LSTM с входными и скрытыми размерами.
        German: Initialisiert die LSTM-Zelle mit Eingangs- und versteckten Dimensionen.
        Chinese: 初始化LSTM单元，输入维度和隐藏状态维度.

        Parameters:
        input_dim (int): English: Input dimension. Russian: Входное измерение. German: Eingabedimension. Chinese: 输入维度.
        hidden_dim (int): English: Hidden state dimension. Russian: Измерение скрытого состояния. German: Versteckte Dimension. Chinese: 隐藏状态维度.
        """
        self.input_dim = input_dim
        self.hidden_dim = hidden_dim

        # Initialize LSTM weights
        self.W_f = np.random.randn(hidden_dim, input_dim)
        self.U_f = np.random.randn(hidden_dim, hidden_dim)
        self.b_f = np.zeros((hidden_dim, 1))

        self.W_i = np.random.randn(hidden_dim, input_dim)
        self.U_i = np.random.randn(hidden_dim, hidden_dim)
        self.b_i = np.zeros((hidden_dim, 1))

        self.W_c = np.random.randn(hidden_dim, input_dim)
        self.U_c = np.random.randn(hidden_dim, hidden_dim)
        self.b_c = np.zeros((hidden_dim, 1))

        self.W_o = np.random.randn(hidden_dim, input_dim)
        self.U_o = np.random.randn(hidden_dim, hidden_dim)
        self.b_o = np.zeros((hidden_dim, 1))

    def sigmoid(self, x):
        """
        English: Apply the sigmoid function to the input.
        Russian: Применить функцию сигмоиды к входу.
        German: Wendet die Sigmoidfunktion auf die Eingabe an.
        Chinese: 对输入应用sigmoid函数.
        """
        return 1 / (1 + np.exp(-x))

    def tanh(self, x):
        """
        English: Apply the tanh function to the input.
        Russian: Применить функцию tanh к входу.
        German: Wendet die Tanh-Funktion auf die Eingabe an.
        Chinese: 对输入应用tanh函数.
        """
        return np.tanh(x)

    def forward(self, x, h_prev, c_prev):
        """
        English: Forward pass for LSTM cell.
        Russian: Прямой проход для ячейки LSTM.
        German: Vorwärtsdurchlauf für die LSTM-Zelle.
        Chinese: LSTM单元的前向传播.

        Parameters:
        x (numpy.ndarray): Input vector.
        h_prev (numpy.ndarray): Previous hidden state.
        c_prev (numpy.ndarray): Previous cell state.

        Returns:
        h_t (numpy.ndarray): Current hidden state.
        c_t (numpy.ndarray): Current cell state.
        """
        f_t = self.sigmoid(np.dot(self.W_f, x) + np.dot(self.U_f, h_prev) + self.b_f)
        i_t = self.sigmoid(np.dot(self.W_i, x) + np.dot(self.U_i, h_prev) + self.b_i)
        c_tilde = self.tanh(np.dot(self.W_c, x) + np.dot(self.U_c, h_prev) + self.b_c)
        c_t = f_t * c_prev + i_t * c_tilde
        o_t = self.sigmoid(np.dot(self.W_o, x) + np.dot(self.U_o, h_prev) + self.b_o)
        h_t = o_t * self.tanh(c_t)

        return h_t, c_t


# SelfAttention
# English: This class implements the self-attention mechanism for capturing dependencies in sequences.
# Russian: Этот класс реализует механизм самовнимания для захвата зависимостей в последовательностях.
# German: Diese Klasse implementiert den Self-Attention-Mechanismus zur Erfassung von Abhängigkeiten in Sequenzen.
# Chinese: 该类实现了用于捕捉序列依赖关系的自注意力机制.
class SelfAttention:
    def __init__(self, embed_dim):
        """
        English: Initialize the self-attention mechanism.
        Russian: Инициализация механизма самовнимания.
        German: Initialisierung des Self-Attention-Mechanismus.
        Chinese: 初始化自注意力机制.

        Parameters:
        embed_dim (int): Embedding dimension.
        """
        self.embed_dim = embed_dim
        self.W_q = np.random.randn(embed_dim, embed_dim)
        self.W_k = np.random.randn(embed_dim, embed_dim)
        self.W_v = np.random.randn(embed_dim, embed_dim)

    def softmax(self, x):
        """
        English: Apply the softmax function to the input.
        Russian: Применить функцию softmax к входу.
        German: Wendet die Softmax-Funktion auf die Eingabe an.
        Chinese: 对输入应用softmax函数.
        """
        e_x = np.exp(x - np.max(x, axis=-1, keepdims=True))
        return e_x / e_x.sum(axis=-1, keepdims=True)

    def forward(self, x):
        """
        English: Forward pass for self-attention.
        Russian: Прямой проход для самовнимания.
        German: Vorwärtsdurchlauf für Self-Attention.
        Chinese: 自注意力的前向传播.

        Parameters:
        x (numpy.ndarray): Input vector (sequence).

        Returns:
        output (numpy.ndarray): Attention-weighted output.
        """
        if len(x.shape) == 3 and x.shape[-1] == 1:
            x = x.squeeze(-1)

        Q = np.dot(x, self.W_q)
        K = np.dot(x, self.W_k)
        V = np.dot(x, self.W_v)

        attention_scores = np.dot(Q, K.T) / np.sqrt(self.embed_dim)
        attention_weights = self.softmax(attention_scores)
        output = np.dot(attention_weights, V)
        return output


# MultiHeadAttention
# English: This class implements multi-head attention for capturing multiple aspects of attention.
# Russian: Этот класс реализует многоголовое внимание для захвата различных аспектов внимания.
# German: Diese Klasse implementiert Multi-Head-Attention, um verschiedene Aspekte der Aufmerksamkeit zu erfassen.
# Chinese: 该类实现了多头注意力机制，用于捕捉多个注意力方面.
class MultiHeadAttention:
    def __init__(self, embed_dim, num_heads):
        """
        English: Initialize multi-head attention.
        Russian: Инициализация многоголового внимания.
        German: Initialisierung von Multi-Head-Attention.
        Chinese: 初始化多头注意力机制.

        Parameters:
        embed_dim (int): Embedding dimension.
        num_heads (int): Number of attention heads.
        """
        self.num_heads = num_heads
        self.attention_heads = [SelfAttention(embed_dim) for _ in range(num_heads)]

    def concatenate(self, head_outputs):
        """
        English: Concatenate outputs from all attention heads.
        Russian: Конкатенация выходов всех голов внимания.
        German: Verkettet Ausgaben von allen Attention-Heads.
        Chinese: 将所有注意力头的输出拼接.

        Parameters:
        head_outputs (list): List of outputs from attention heads.

        Returns:
        combined_output (numpy.ndarray): Concatenated output.
        """
        return np.concatenate(head_outputs, axis=-1)

    def forward(self, x):
        """
        English: Forward pass for multi-head attention.
        Russian: Прямой проход для многоголового внимания.
        German: Vorwärtsdurchlauf für Multi-Head-Attention.
        Chinese: 多头注意力的前向传播.

        Parameters:
        x (numpy.ndarray): Input vector (sequence).

        Returns:
        combined_output (numpy.ndarray): Combined output from all attention heads.
        """
        head_outputs = [head.forward(x) for head in self.attention_heads]
        combined_output = self.concatenate(head_outputs)
        return combined_output


# ResidualSelfAttention
# English: This class implements residual connections for self-attention layers.
# Russian: Этот класс реализует остаточные связи для слоев самовнимания.
# German: Diese Klasse implementiert Restverbindungen für Self-Attention-Schichten.
# Chinese: 该类实现了自注意力层的残差连接.
class ResidualSelfAttention(SelfAttention):
    def forward(self, x):
        """
        English: Forward pass with residual connections.
        Russian: Прямой проход с остаточными связями.
        German: Vorwärtsdurchlauf mit Restverbindungen.
        Chinese: 带残差连接的前向传播.

        Parameters:
        x (numpy.ndarray): Input vector.

        Returns:
        output (numpy.ndarray): Residual-connected output.
        """
        attention_output = super().forward(x)
        return x + attention_output

# RegularizedAttention
# English: This class adds regularization to the self-attention mechanism.
# Russian: Этот класс добавляет регуляризацию к механизму самовнимания.
# German: Diese Klasse fügt dem Self-Attention-Mechanismus Regularisierung hinzu.
# Chinese: 该类为自注意力机制添加了正则化.
class RegularizedAttention(SelfAttention):
    def __init__(self, embed_dim, reg_lambda=0.01):
        """
        English: Initialize regularized self-attention with a regularization factor.
        Russian: Инициализация регуляризованного самовнимания с регуляризационным коэффициентом.
        German: Initialisiert regulierte Self-Attention mit einem Regularisierungsfaktor.
        Chinese: 用正则化因子初始化正则化自注意力.

        Parameters:
        embed_dim (int): Embedding dimension.
        reg_lambda (float): Regularization factor.
        """
        super().__init__(embed_dim)
        self.reg_lambda = reg_lambda

    def regularize(self, attention_weights):
        """
        English: Apply regularization to attention weights.
        Russian: Применить регуляризацию к весам внимания.
        German: Regularisierung auf die Attention-Gewichte anwenden.
        Chinese: 对注意力权重应用正则化.

        Parameters:
        attention_weights (numpy.ndarray): Attention weights.

        Returns:
        reg_loss (float): Regularization loss.
        """
        return self.reg_lambda * np.sum(np.square(attention_weights))

    def forward(self, x):
        """
        English: Forward pass for regularized attention.
        Russian: Прямой проход для регуляризованного внимания.
        German: Vorwärtsdurchlauf für regulierte Aufmerksamkeit.
        Chinese: 正则化注意力的前向传播.

        Parameters:
        x (numpy.ndarray): Input vector.

        Returns:
        output (numpy.ndarray): Attention-weighted output.
        reg_loss (float): Regularization loss.
        """
        if len(x.shape) == 3 and x.shape[-1] == 1:
            x = x.squeeze(-1)

        Q = np.dot(x, self.W_q)
        K = np.dot(x, self.W_k)
        V = np.dot(x, self.W_v)

        attention_scores = np.dot(Q, K.T) / np.sqrt(self.embed_dim)
        attention_weights = self.softmax(attention_scores)

        reg_loss = self.regularize(attention_weights)
        output = np.dot(attention_weights, V)
        return output, reg_loss


# HybridModelWithImprovements
# English: This class combines LSTM and attention mechanisms with residual connections and regularization.
# Russian: Этот класс объединяет LSTM и механизмы внимания с остаточными связями и регуляризацией.
# German: Diese Klasse kombiniert LSTM- und Attention-Mechanismen mit Restverbindungen und Regularisierung.
# Chinese: 该类结合了LSTM与注意力机制，并添加了残差连接和正则化.
class HybridModelWithImprovements:
    def __init__(self, input_dim, hidden_dim, embed_dim, num_heads, output_dim, reg_lambda=0.01):
        """
        English: Initialize the hybrid model with LSTM, attention, and regularization.
        Russian: Инициализация гибридной модели с LSTM, вниманием и регуляризацией.
        German: Initialisiert das hybride Modell mit LSTM, Attention und Regularisierung.
        Chinese: 用LSTM、注意力机制和正则化初始化混合模型.

        Parameters:
        input_dim (int): Input dimension.
        hidden_dim (int): LSTM hidden state dimension.
        embed_dim (int): Attention embedding dimension.
        num_heads (int): Number of attention heads.
        output_dim (int): Output dimension.
        reg_lambda (float): Regularization factor.
        """
        self.lstm = LSTMCell(input_dim, hidden_dim)
        self.multi_head_attention = MultiHeadAttention(embed_dim, num_heads)
        self.regularized_attention = RegularizedAttention(embed_dim, reg_lambda)
        self.output_dim = output_dim
        self.W_out = np.random.randn(output_dim, hidden_dim + embed_dim * num_heads)
        self.b_out = np.zeros((output_dim, 1))

    def forward(self, x_seq):
        """
        English: Forward pass for the hybrid model.
        Russian: Прямой проход для гибридной модели.
        German: Vorwärtsdurchlauf für das hybride Modell.
        Chinese: 混合模型的前向传播.

        Parameters:
        x_seq (numpy.ndarray): Input sequence.

        Returns:
        y_pred (numpy.ndarray): Model predictions.
        reg_loss (float): Regularization loss.
        """
        h_prev = np.zeros((self.lstm.hidden_dim, 1))
        c_prev = np.zeros((self.lstm.hidden_dim, 1))

        lstm_outputs = []
        for x in x_seq:
            h_prev, c_prev = self.lstm.forward(x, h_prev, c_prev)
            lstm_outputs.append(h_prev)

        lstm_outputs = np.stack(lstm_outputs, axis=0)

        # Attention mechanism
        attention_output = self.multi_head_attention.forward(lstm_outputs)

        # Use last step of LSTM and attention
        attention_output_last = attention_output[-1]
        lstm_output_last = lstm_outputs[-1]

        if attention_output_last.ndim == 1:
            attention_output_last = np.expand_dims(attention_output_last, axis=-1)

        if lstm_output_last.ndim == 1:
            lstm_output_last = np.expand_dims(lstm_output_last, axis=-1)

        combined_output = np.concatenate((lstm_output_last, attention_output_last), axis=0)
        combined_output = combined_output.reshape(-1, 1)

        y_pred = np.dot(self.W_out, combined_output) + self.b_out
        reg_loss = 0

        return y_pred, reg_loss


# Metrics for model evaluation
# English: Provides functions for evaluating the model's performance.
# Russian: Содержит функции для оценки производительности модели.
# German: Bietet Funktionen zur Bewertung der Modellleistung.
# Chinese: 提供评估模型性能的功能.
def mean_squared_error(y_true, y_pred):
    """
    English: Calculate mean squared error (MSE).
    Russian: Вычислить среднеквадратичную ошибку (MSE).
    German: Berechnet den mittleren quadratischen Fehler (MSE).
    Chinese: 计算均方误差 (MSE).

    Parameters:
    y_true (numpy.ndarray): True values.
    y_pred (numpy.ndarray): Predicted values.

    Returns:
    mse (float): Mean squared error.
    """
    return np.mean((y_true - y_pred) ** 2)


def r_squared(y_true, y_pred):
    """
    English: Calculate R-squared metric.
    Russian: Вычислить коэффициент детерминации (R^2).
    German: Berechnet den R-Quadrat-Wert.
    Chinese: 计算决定系数 (R^2).

    Parameters:
    y_true (numpy.ndarray): True values.
    y_pred (numpy.ndarray): Predicted values.

    Returns:
    r2 (float): R-squared value.
    """
    ss_res = np.sum((y_true - y_pred) ** 2)
    ss_tot = np.sum((y_true - np.mean(y_true)) ** 2)

    if ss_tot == 0:
        return 0  # Return 0 when y_true has no variance.

    return 1 - (ss_res / ss_tot)


def mean_absolute_error(y_true, y_pred):
    """
    English: Calculate mean absolute error (MAE).
    Russian: Вычислить среднюю абсолютную ошибку (MAE).
    German: Berechnet den mittleren absoluten Fehler (MAE).
    Chinese: 计算平均绝对误差 (MAE).

    Parameters:
    y_true (numpy.ndarray): True values.
    y_pred (numpy.ndarray): Predicted values.

    Returns:
    mae (float): Mean absolute error.
    """
    return np.mean(np.abs(y_true - y_pred))


def evaluate_model(y_true, y_pred):
    """
    English: Evaluate the model using MSE, R^2, and MAE.
    Russian: Оценка модели с использованием MSE, R^2 и MAE.
    German: Bewertung des Modells mit MSE, R^2 und MAE.
    Chinese: 使用MSE、R^2和MAE评估模型.

    Parameters:
    y_true (numpy.ndarray): True values.
    y_pred (numpy.ndarray): Predicted values.

    Returns:
    metrics (dict): Dictionary with MSE, R^2, and MAE.
    """
    mse = mean_squared_error(y_true, y_pred)
    r2 = r_squared(y_true, y_pred)
    mae = mean_absolute_error(y_true, y_pred)
    return {"MSE": mse, "R^2": r2, "MAE": mae}


# save_plot
# English: This function saves a plot of data.
# Russian: Эта функция сохраняет график данных.
# German: Diese Funktion speichert ein Diagramm der Daten.
# Chinese: 此函数保存数据的图表.
def save_plot(x, y, title, xlabel, ylabel, filename, save_dir="reports/"):
    """
    English: Save a plot to a file.

    Russian: Сохранить график в файл.

    German: Speichert ein Diagramm in einer Datei.

    Chinese: 将图表保存到文件中.

    Parameters:
    x (list or numpy.ndarray): X-axis data.
    y (list or numpy.ndarray): Y-axis data.
    title (str): Plot title.
    xlabel (str): X-axis label.
    ylabel (str): Y-axis label.
    filename (str): File name to save the plot.
    save_dir (str): Directory to save the file. Default is "reports/".
    """
    plt.figure()
    plt.plot(x, y)
    plt.title(title)
    plt.xlabel(xlabel)
    plt.ylabel(ylabel)
    plt.grid(True)
    if not os.path.exists(save_dir):
        os.makedirs(save_dir)
    plt.savefig(os.path.join(save_dir, filename))
    plt.close()

# plot_losses
# English: This function saves training and validation loss plots.
# Russian: Эта функция сохраняет графики потерь при обучении и валидации.
# German: Diese Funktion speichert Diagramme für Trainings- und Validierungsverlust.
# Chinese: 此函数保存训练损失和验证损失的图表.
def plot_losses(train_loss, val_loss, save_dir="reports/"):
    """
    English: Save plots for training and validation loss.

    Russian: Сохранить графики потерь при обучении и валидации.

    German: Speichert Diagramme für Trainings- und Validierungsverlust.

    Chinese: 保存训练和验证损失的图表.

    Parameters:
    train_loss (list): List of training loss values.
    val_loss (list): List of validation loss values.
    save_dir (str): Directory to save the plots. Default is "reports/".
    """
    epochs = list(range(1, len(train_loss) + 1))
    save_plot(epochs, train_loss, "Training Loss", "Epochs", "Loss", "train_loss.png", save_dir)
    save_plot(epochs, val_loss, "Validation Loss", "Epochs", "Loss", "val_loss.png", save_dir)

# generate_word_report
# English: This function generates a Word report with model metrics and plots.
# Russian: Эта функция генерирует отчет в Word с метриками модели и графиками.
# German: Diese Funktion erstellt einen Word-Bericht mit Modellmetriken und Diagrammen.
# Chinese: 此函数生成包含模型指标和图表的Word报告.
def generate_word_report(report_data, save_dir="reports/", filename="report.docx"):
    """
    English: Generate a Word report with metrics and plots.

    Russian: Сгенерировать отчет в Word с метриками и графиками.

    German: Erstellen Sie einen Word-Bericht mit Metriken und Diagrammen.

    Chinese: 生成包含指标和图表的Word报告.

    Parameters:
    report_data (dict): Report data containing metrics and plot filenames.
    save_dir (str): Directory to save the report. Default is "reports/".
    filename (str): Filename for the report. Default is "report.docx".
    """
    doc = Document()
    doc.add_heading('Model Report', 0)
    doc.add_paragraph('This report contains data visualization, metrics, and results of the model.')
    doc.add_heading('Metrics', level=1)
    for metric, value in report_data["metrics"].items():
        doc.add_paragraph(f"{metric}: {value:.4f}")
    doc.add_heading('Plots', level=1)
    for plot_name in report_data["plots"]:
        doc.add_paragraph(plot_name)
        doc.add_picture(os.path.join(save_dir, plot_name))
    if not os.path.exists(save_dir):
        os.makedirs(save_dir)
    doc.save(os.path.join(save_dir, filename))

# generate_excel_report
# English: This function generates an Excel report with model metrics.
# Russian: Эта функция генерирует отчет в Excel с метриками модели.
# German: Diese Funktion erstellt einen Excel-Bericht mit Modellmetriken.
# Chinese: 此函数生成包含模型指标的Excel报告.
def generate_excel_report(report_data, save_dir="reports/", filename="report.xlsx"):
    """
    English: Generate an Excel report with metrics.

    Russian: Сгенерировать отчет в Excel с метриками.

    German: Erstellen Sie einen Excel-Bericht mit Metriken.

    Chinese: 生成包含指标的Excel报告.

    Parameters:
    report_data (dict): Report data containing metrics.
    save_dir (str): Directory to save the report. Default is "reports/".
    filename (str): Filename for the report. Default is "report.xlsx".
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Metrics"
    ws.append(["Metric", "Value"])
    for metric, value in report_data["metrics"].items():
        ws.append([metric, value])
    if not os.path.exists(save_dir):
        os.makedirs(save_dir)
    wb.save(os.path.join(save_dir, filename))

# run_model_and_generate_report
# English: This function runs the model, generates predictions, and creates reports.
# Russian: Эта функция запускает модель, генерирует предсказания и создает отчеты.
# German: Diese Funktion führt das Modell aus, erstellt Vorhersagen und generiert Berichte.
# Chinese: 此函数运行模型、生成预测并创建报告.
def run_model_and_generate_report(x_seq, y_true, model, save_dir="reports/"):
    """
    English: Run the model, generate predictions, and create Word and Excel reports.

    Russian: Запустить модель, сгенерировать предсказания и создать отчеты в Word и Excel.

    German: Führen Sie das Modell aus, erstellen Sie Vorhersagen und erstellen Sie Word- und Excel-Berichte.

    Chinese: 运行模型，生成预测并创建Word和Excel报告.

    Parameters:
    x_seq (numpy.ndarray): Input sequence for the model.
    y_true (numpy.ndarray): True values for evaluation.
    model (object): The model to be run.
    save_dir (str): Directory to save the reports. Default is "reports/".
    """
    y_pred, reg_loss = model.forward(x_seq)
    metrics = evaluate_model(y_true, y_pred)
    train_loss = [0.1, 0.05, 0.03, 0.01]  
    val_loss = [0.12, 0.07, 0.05, 0.02]
    plot_losses(train_loss, val_loss, save_dir)
    report_data = {
        "metrics": metrics,
        "plots": ["train_loss.png", "val_loss.png"]
    }
    generate_word_report(report_data, save_dir)
    generate_excel_report(report_data, save_dir)