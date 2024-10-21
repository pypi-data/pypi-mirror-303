import os
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from causalml.inference.meta import LRSRegressor
import shap
import pandas as pd
from openpyxl import load_workbook
from openpyxl.drawing.image import Image
from docx import Document
from docx.shared import Inches

class IljicevsCausalModel:
    """
    Класс для работы с причинно-следственной моделью.
    """
    def __init__(self, model_type='meta'):
        """
        Инициализация причинно-следственной модели.
        """
        if model_type == 'meta':
            self.model = LRSRegressor()
        elif model_type == 'tree':
            from causalml.inference.tree import CausalTreeRegressor
            self.model = CausalTreeRegressor()

    def fit(self, X, treatment, y):
        """
        Обучение причинно-следственной модели.
        """
        self.model.fit(X, treatment, y)

    def predict(self, X):
        """
        Предсказание с использованием причинно-следственной модели.
        """
        return self.model.predict(X)

    def feature_importance(self, output_dir):
        """
        Визуализация важности признаков на основе причинно-следственного анализа.
        """
        params = None 
        try:
            if 1 in self.model.models:
                internal_model = self.model.models[1] 
                if hasattr(internal_model, 'coefficients'):
                    params = internal_model.coefficients 
                    print("Коэффициенты модели:", params)
                else:
                    print("Модель не содержит атрибута 'coefficients'.")
                    return None

            else:
                print("Модель для ключа 1 не найдена.")
                return None

        except AttributeError as e:
            print(f"Не удалось извлечь коэффициенты. Ошибка: {e}")
            return None
        except Exception as e:
            print(f"Произошла ошибка: {str(e)}")
            return None

        if params is not None:
            plt.figure(figsize=(10, 6))
            sns.barplot(x=np.arange(len(params)), y=params)
            plt.xlabel("Признаки")
            plt.ylabel("Важность (коэффициенты)")
            plt.title("Важность признаков для причинно-следственной модели")

            if not os.path.exists(output_dir):
                os.makedirs(output_dir)

            output_path = os.path.join(output_dir, "causal_feature_importance.png")
            plt.savefig(output_path)
            plt.close()
            print(f"Визуализация важности признаков сохранена в: {output_path}")
            return output_path
        else:
            print("Не удалось получить важность признаков.")
            return None


    def plot_causal_effects(self, X, treatment, y, output_dir):
        """
        Визуализация причинно-следственных эффектов.
        """
        uplift = self.model.predict(X)
        plt.figure(figsize=(10, 6))
        plt.scatter(treatment, uplift, c=y, cmap='bwr', edgecolors='k')
        plt.xlabel("Лечение (Воздействие)")
        plt.ylabel("Причинно-следственный эффект")
        plt.title("Причинно-следственный эффект на основе модели")
        
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        output_path = os.path.join(output_dir, "causal_effects.png")
        print(f"Визуализация важности признаков сохранена в: {output_path}")
        plt.savefig(output_path)
        plt.close()
        return output_path
    

    def plot_shap_values(self, X, output_dir):
        """
        Визуализация SHAP-значений для объяснения модели.
        """
        explainer = shap.Explainer(self.model.predict, X)
        shap_values = explainer(X)

        plt.figure(figsize=(10, 6))
        shap.summary_plot(shap_values, X, show=False)  # Отключаем вывод графика
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        output_path = os.path.join(output_dir, "shap_summary_plot.png")
        plt.savefig(output_path, bbox_inches='tight')
        plt.close()
        print(f"Визуализация SHAP-значений сохранена в: {output_path}")
        return output_path

    def counterfactual_analysis(self, X, treatment, feature_index, new_value):
        """
        Контрфактический анализ: что бы произошло, если бы один из признаков изменился.
        """
        X_new = X.copy()
        X_new[:, feature_index] = new_value  # Изменяем значение признака
        
        # Предсказываем исход с новыми данными
        uplift_original = self.model.predict(X)
        uplift_new = self.model.predict(X_new)
        
        delta_uplift = uplift_new - uplift_original
        return delta_uplift

    def estimate_confidence_intervals(self, X, treatment, n_bootstrap=1000, alpha=0.05):
        """
        Оценка доверительных интервалов для предсказаний с использованием бутстрапа.
        """
        predictions = []
        
        for _ in range(n_bootstrap):
            idx = np.random.choice(len(X), len(X), replace=True)
            X_sample = X[idx]
            treatment_sample = treatment[idx]
            pred_sample = self.model.predict(X_sample)
            predictions.append(pred_sample)
        
        predictions = np.array(predictions)
        
        lower_bound = np.percentile(predictions, alpha/2*100, axis=0)
        upper_bound = np.percentile(predictions, (1-alpha/2)*100, axis=0)
        
        return lower_bound, upper_bound

    def analyze_interactions(self, X, treatment, output_dir):
        """
        Анализ взаимодействий между признаками с выбором правильного объяснителя в зависимости от типа модели.
        """
        # Проверка типа модели для выбора правильного SHAP Explainer
        model_type = type(self.model)

        if model_type.__name__ == 'CausalTreeRegressor':
            # Для моделей на основе деревьев решений используем TreeExplainer
            explainer = shap.TreeExplainer(self.model)
            shap_interaction_values = explainer.shap_interaction_values(X)
        elif model_type.__name__ == 'LRSRegressor':
            # Для линейных моделей или обобщенных моделей используем обычный Explainer
            explainer = shap.Explainer(self.model.predict, X)
            shap_interaction_values = explainer(X)
        else:
            print(f"Тип модели '{model_type.__name__}' не поддерживается для анализа взаимодействий.")
            return None

        # Визуализация результатов
        plt.figure(figsize=(10, 6))
        shap.summary_plot(shap_interaction_values, X, show=False)
        plt.title("Взаимодействие признаков")

        # Сохранение графика
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        output_path = os.path.join(output_dir, "interaction_plot.png")
        plt.savefig(output_path, bbox_inches='tight')
        plt.close()
        print(f"Взаимодействие признаков сохранено в: {output_path}")

        return output_path
    
    def generate_report(self, X_test, treatment_test, y_test, output_dir):
        """
        Генерация отчета в формате Excel и Word, содержащего все важные результаты и визуализации.
        """
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        feature_importance_path = self.feature_importance(output_dir)
        causal_effects_path = self.plot_causal_effects(X_test, treatment_test, y_test, output_dir)
        shap_path = self.plot_shap_values(X_test, output_dir)
        interaction_path = self.analyze_interactions(X_test, treatment_test, output_dir)

        excel_report_path = os.path.join(output_dir, "causal_analysis_report.xlsx")

        # Coefficients and feature importances
        if hasattr(self.model.models[1], 'coefficients'):
            coefficients = self.model.models[1].coefficients
            df_importance = pd.DataFrame({
                'Feature': [f'Feature_{i}' for i in range(len(coefficients))],
                'Coefficient': coefficients
            })
            with pd.ExcelWriter(excel_report_path, engine='openpyxl') as writer:
                df_importance.to_excel(writer, sheet_name='Feature Importance', index=False)

                # Add the counterfactual analysis result
                delta_uplift = self.counterfactual_analysis(X_test, treatment_test, feature_index=0, new_value=1.0)
                if delta_uplift.ndim > 1:
                    delta_uplift = delta_uplift.flatten()

                df_counterfactual = pd.DataFrame({
                    'Observation': np.arange(len(delta_uplift)),
                    'Delta Uplift': delta_uplift
                })
                df_counterfactual.to_excel(writer, sheet_name='Counterfactual Analysis', index=False)

                # Add confidence intervals
                lower_bound, upper_bound = self.estimate_confidence_intervals(X_test, treatment_test)
                df_confidence_intervals = pd.DataFrame({
                    'Lower Bound': lower_bound.flatten(),
                    'Upper Bound': upper_bound.flatten()
                })
                df_confidence_intervals.to_excel(writer, sheet_name='Confidence Intervals', index=False)

                # Save paths to visualizations in Excel
                df_visualization_paths = pd.DataFrame({
                    'Visualization': ['Feature Importance', 'Causal Effects', 'SHAP Summary', 'Interactions'],
                    'File Path': [feature_importance_path, causal_effects_path, shap_path, interaction_path]
                })
                df_visualization_paths.to_excel(writer, sheet_name='Visualizations', index=False)

        wb = load_workbook(excel_report_path)
        sheet_name = 'Graphs'
        if sheet_name not in wb.sheetnames:
            ws = wb.create_sheet(sheet_name)
        else:
            ws = wb[sheet_name]

        img_feature_importance = Image(feature_importance_path)
        img_causal_effects = Image(causal_effects_path)
        img_shap = Image(shap_path)
        img_interaction = Image(interaction_path)

        ws.add_image(img_feature_importance, 'A1')
        ws.add_image(img_causal_effects, 'A20')
        ws.add_image(img_shap, 'A40')
        ws.add_image(img_interaction, 'A60')

        wb.save(excel_report_path)

        word_report_path = os.path.join(output_dir, "causal_analysis_report.docx")
        doc = Document()

        doc.add_heading('Causal Analysis Report', 0)
        doc.add_heading('Feature Importance', level=1)
        doc.add_paragraph('Feature importance calculated from the causal model:')

        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Feature'
        hdr_cells[1].text = 'Coefficient'
        for i, coef in enumerate(coefficients):
            row_cells = table.add_row().cells
            row_cells[0].text = f'Feature_{i}'
            row_cells[1].text = str(coef)

        doc.add_heading('Visualizations', level=1)
        doc.add_paragraph('Below are visualizations generated by the model:')

        doc.add_paragraph('Feature Importance:')
        doc.add_picture(feature_importance_path, width=Inches(4))

        doc.add_paragraph('Causal Effects:')
        doc.add_picture(causal_effects_path, width=Inches(4))

        doc.add_paragraph('SHAP Summary:')
        doc.add_picture(shap_path, width=Inches(4))

        doc.add_paragraph('Feature Interactions:')
        doc.add_picture(interaction_path, width=Inches(4))

        doc.add_heading('Counterfactual Analysis', level=1)
        doc.add_paragraph('Changes in predictions when the first feature is changed to 1.0:')
        
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Observation'
        hdr_cells[1].text = 'Delta Uplift'
        for i, delta in enumerate(delta_uplift):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = str(delta)

        doc.add_heading('Confidence Intervals', level=1)
        doc.add_paragraph('Confidence intervals calculated for the predictions:')

        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Lower Bound'
        hdr_cells[1].text = 'Upper Bound'
        for lb, ub in zip(lower_bound, upper_bound):
            row_cells = table.add_row().cells
            row_cells[0].text = str(lb)
            row_cells[1].text = str(ub)

        doc.save(word_report_path)
        print(f"Отчет в формате Word сохранен в: {word_report_path}")

        return excel_report_path, word_report_path
