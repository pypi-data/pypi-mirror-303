import torch
import torch.nn as nn
import torch.nn.functional as F
import matplotlib.pyplot as plt
from torch.utils.tensorboard import SummaryWriter
from docx import Document
from docx.shared import Inches

import os
os.makedirs("/logs", exist_ok=True)
writer = SummaryWriter(log_dir="/logs")

class AdaptiveActivation(nn.Module):
    def __init__(self, num_neurons, num_functions=5):
        super(AdaptiveActivation, self).__init__()
        self.weights = nn.Parameter(torch.randn(num_neurons, num_functions, requires_grad=True))
        self.attention_weights = nn.Parameter(torch.ones(num_neurons, requires_grad=True)) 
        self.activations = [F.relu, torch.sigmoid, torch.tanh, F.elu, lambda x: x * torch.sigmoid(x)]
    
    def forward(self, x):
        normalized_weights = F.softmax(self.weights, dim=1)
        attention = F.sigmoid(self.attention_weights).unsqueeze(0).expand(x.size(0), -1)
        
        output = torch.zeros_like(x)
        for i, activation in enumerate(self.activations):
            activation_result = activation(x)
            weight = normalized_weights[:, i].unsqueeze(0).expand_as(activation_result)
            attention_expanded = attention.expand_as(activation_result)
            
            output += weight * activation_result * attention_expanded
        
        return output


class DropConnectActivation(AdaptiveActivation):
    def forward(self, x, drop_prob=0.5):
        normalized_weights = F.softmax(self.weights, dim=1)
        output = torch.zeros_like(x)
        for i, activation in enumerate(self.activations):
            mask = (torch.rand_like(x) > drop_prob).float()
            output += normalized_weights[:, i].unsqueeze(1) * activation(x) * mask
        return output


class AdaptiveLayer(nn.Module):
    def __init__(self, input_size, output_size, drop_connect=False):
        super(AdaptiveLayer, self).__init__()
        self.fc = nn.Linear(input_size, output_size)
        if drop_connect:
            self.adaptive_activation = DropConnectActivation(output_size)
        else:
            self.adaptive_activation = AdaptiveActivation(output_size)

    def forward(self, x):
        x = self.fc(x)
        x = self.adaptive_activation(x)
        return x


class MultiTaskLearningNet(nn.Module):
    def __init__(self, input_size, hidden_size, output_size_regression, output_size_classification):
        super(MultiTaskLearningNet, self).__init__()
        self.layer1 = AdaptiveLayer(input_size, hidden_size)
        self.regression_head = nn.Linear(hidden_size, output_size_regression)
        self.classification_head = nn.Linear(hidden_size, output_size_classification)
        self.loss_history = [] 

    def forward(self, x):
        x = self.layer1(x)
        regression_output = self.regression_head(x)
        classification_output = self.classification_head(x)
        return regression_output, classification_output

    def log_metrics(self, loss_regression, loss_classification, epoch):
        total_loss = loss_regression + loss_classification
        self.loss_history.append(total_loss.item()) 
        writer.add_scalar('Loss/Regression', loss_regression, epoch)
        writer.add_scalar('Loss/Classification', loss_classification, epoch)

    def log_activation_weights(self, epoch):
        for name, param in self.named_parameters():
            if "weights" in name:
                writer.add_histogram(f'Weights/{name}', param, epoch)

    def regularization_loss(self):
        reg_loss = 0.0
        for param in self.parameters():
            if param.requires_grad:
                reg_loss += torch.sum(torch.abs(param)) 
        return reg_loss

def plot_loss_history(loss_history):
    plt.plot(range(len(loss_history)), loss_history, label='Loss')
    plt.title('Loss History')
    plt.xlabel('Epoch')
    plt.ylabel('Loss')
    plt.legend()
    plt.show()

def save_loss_plot(loss_history, filename="loss_plot.png"):
    plt.figure()
    plt.plot(range(len(loss_history)), loss_history, label='Loss')
    plt.title('Loss History')
    plt.xlabel('Epoch')
    plt.ylabel('Loss')
    plt.legend()
    plt.savefig(filename) 
    plt.close()

def generate_report(model, loss_history, report_filename="training_report.docx"):
    doc = Document()
    doc.add_heading('Training Report', 0)
    
    doc.add_heading('Model Description', level=1)
    model_description = f"The model has {sum(p.numel() for p in model.parameters() if p.requires_grad)} trainable parameters."
    doc.add_paragraph(model_description)
    
    doc.add_heading('Training Metrics', level=1)
    final_loss = loss_history[-1]
    doc.add_paragraph(f"Final Loss: {final_loss}")
    
    plot_filename = "loss_plot.png"
    save_loss_plot(loss_history, filename=plot_filename)
    
    doc.add_heading('Loss History', level=1)
    doc.add_paragraph('The following plot shows the loss history over epochs:')
    doc.add_picture(plot_filename, width=Inches(5.0))
    
    report_path = os.path.join(os.getcwd(), report_filename)
    doc.save(report_path)
    print(f"Report saved to {report_path}")