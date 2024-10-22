import numpy as np
import matplotlib.pyplot as plt
import os
import pandas as pd
from sklearn.metrics import accuracy_score, f1_score, confusion_matrix
from sklearn.model_selection import train_test_split
from docx import Document
from docx.shared import Inches
import os

class ProgressiveModel:
    def __init__(self, input_size, hidden_size_1, hidden_size_2, output_size, l2_reg=0.0, l1_reg=0.0, dropout_rate=0.0):
        self.weights_input_hidden1 = np.random.randn(input_size, hidden_size_1) * np.sqrt(2. / input_size)
        self.weights_hidden1_hidden2 = np.random.randn(hidden_size_1, hidden_size_2) * np.sqrt(2. / hidden_size_1)
        self.weights_hidden2_output = np.random.randn(hidden_size_2, output_size) * np.sqrt(2. / hidden_size_2)
        self.bias_hidden1 = np.zeros((1, hidden_size_1))
        self.bias_hidden2 = np.zeros((1, hidden_size_2))
        self.bias_output = np.zeros((1, output_size))
        self.l2_reg = l2_reg
        self.l1_reg = l1_reg
        self.dropout_rate = dropout_rate
        self.hidden_dropout_mask = None
    
    def relu(self, x):
        return np.maximum(0, x)
    
    def softmax(self, x):
        exp_x = np.exp(x - np.max(x, axis=1, keepdims=True))
        return exp_x / np.sum(exp_x, axis=1, keepdims=True)
    
    def forward(self, x, train=True):
        self.z_hidden1 = np.dot(x, self.weights_input_hidden1) + self.bias_hidden1
        self.a_hidden1 = self.relu(self.z_hidden1)
        
        self.z_hidden2 = np.dot(self.a_hidden1, self.weights_hidden1_hidden2) + self.bias_hidden2
        self.a_hidden2 = self.relu(self.z_hidden2)
        
        if train and self.dropout_rate > 0:
            self.hidden_dropout_mask = np.random.rand(*self.a_hidden2.shape) > self.dropout_rate
            self.a_hidden2 *= self.hidden_dropout_mask
        
        self.z_output = np.dot(self.a_hidden2, self.weights_hidden2_output) + self.bias_output
        self.a_output = self.softmax(self.z_output)
        return self.a_output
    
    def backward(self, x, y, learning_rate=0.01):
        m = y.shape[0]
        dz_output = self.a_output - y
        dw_hidden2_output = np.dot(self.a_hidden2.T, dz_output) / m + self.l2_reg * self.weights_hidden2_output + self.l1_reg * np.sign(self.weights_hidden2_output)
        db_output = np.sum(dz_output, axis=0, keepdims=True) / m

        dz_hidden2 = np.dot(dz_output, self.weights_hidden2_output.T) * (self.z_hidden2 > 0)
        dw_hidden1_hidden2 = np.dot(self.a_hidden1.T, dz_hidden2) / m + self.l2_reg * self.weights_hidden1_hidden2 + self.l1_reg * np.sign(self.weights_hidden1_hidden2)
        db_hidden2 = np.sum(dz_hidden2, axis=0, keepdims=True) / m
        
        dz_hidden1 = np.dot(dz_hidden2, self.weights_hidden1_hidden2.T) * (self.z_hidden1 > 0)
        dw_input_hidden1 = np.dot(x.T, dz_hidden1) / m + self.l2_reg * self.weights_input_hidden1 + self.l1_reg * np.sign(self.weights_input_hidden1)
        db_hidden1 = np.sum(dz_hidden1, axis=0, keepdims=True) / m

        self.weights_hidden2_output -= learning_rate * dw_hidden2_output
        self.bias_output -= learning_rate * db_output
        self.weights_hidden1_hidden2 -= learning_rate * dw_hidden1_hidden2
        self.bias_hidden2 -= learning_rate * db_hidden2
        self.weights_input_hidden1 -= learning_rate * dw_input_hidden1
        self.bias_hidden1 -= learning_rate * db_hidden1

    def predict(self, x):
        probabilities = self.forward(x, train=False)
        return np.argmax(probabilities, axis=1)

# Вспомогательные функции
def one_hot_encode(y, num_classes):
    one_hot = np.zeros((y.size, num_classes))
    one_hot[np.arange(y.size), y] = 1
    return one_hot

def early_stopping_check(accuracies, patience=5):
    if len(accuracies) > patience and all(acc <= accuracies[-patience] for acc in accuracies[-patience:]):
        return True
    return False

def adjust_learning_rate(learning_rate, epoch, decay_rate=0.95):
    return learning_rate * (decay_rate ** epoch)

def generate_new_features(data, max_new_features=100):
    squared_features = data ** 2
    if squared_features.shape[1] > max_new_features:
        squared_features = squared_features[:, :max_new_features]
    return np.hstack([data, squared_features])

def augment_data(data, noise_level=0.05):
    noise = np.random.randn(*data.shape) * noise_level
    return data + noise

def create_mini_batches(data, targets, batch_size):
    indices = np.arange(data.shape[0])
    np.random.shuffle(indices)
    for start_idx in range(0, data.shape[0] - batch_size + 1, batch_size):
        excerpt = indices[start_idx:start_idx + batch_size]
        yield data[excerpt], targets[excerpt]

def save_metric_plots(accuracies, losses, f1_scores, report_dir):
    if not os.path.exists(report_dir):
        os.makedirs(report_dir)

    # Визуализация точности
    plt.figure(figsize=(8, 6))
    plt.plot(accuracies, label='Accuracy', color='b')
    plt.title('Accuracy over epochs')
    plt.xlabel('Epoch')
    plt.ylabel('Accuracy')
    plt.legend()
    plt.savefig(os.path.join(report_dir, "accuracy_plot.png"))
    plt.close()

    # Визуализация потерь (loss)
    plt.figure(figsize=(8, 6))
    plt.plot(losses, label='Loss', color='r')
    plt.title('Loss over epochs')
    plt.xlabel('Epoch')
    plt.ylabel('Loss')
    plt.legend()
    plt.savefig(os.path.join(report_dir, "loss_plot.png"))
    plt.close()

    # Визуализация F1 Score
    plt.figure(figsize=(8, 6))
    plt.plot(f1_scores, label='F1 Score', color='g')
    plt.title('F1 Score over epochs')
    plt.xlabel('Epoch')
    plt.ylabel('F1 Score')
    plt.legend()
    plt.savefig(os.path.join(report_dir, "f1_score_plot.png"))
    plt.close()


def add_text_summary(doc, accuracies, losses, f1_scores):
    doc.add_heading('Textual Summary', level=1)

    max_acc = max(accuracies)
    min_loss = min(losses)
    max_f1 = max(f1_scores)

    doc.add_paragraph(f"Max Accuracy achieved: {max_acc:.2f}")
    doc.add_paragraph(f"Min Loss achieved: {min_loss:.4f}")
    doc.add_paragraph(f"Max F1 Score achieved: {max_f1:.2f}")

    doc.add_paragraph(f"Total epochs run: {len(accuracies)}")
    doc.add_paragraph(f"Early stopping was triggered after {len(accuracies)} epochs.")


def save_confusion_matrix(cm, epoch, report_dir):
    # Убедимся, что директория существует
    if not os.path.exists(report_dir):
        os.makedirs(report_dir)
    
    plt.figure(figsize=(5, 5))
    plt.imshow(cm, interpolation='nearest', cmap=plt.get_cmap('Blues'))
    plt.title(f'Confusion Matrix at Epoch {epoch}')
    plt.colorbar()
    plt.xlabel('Predicted label')
    plt.ylabel('True label')
    tick_marks = np.arange(len(cm))
    plt.xticks(tick_marks)
    plt.yticks(tick_marks)
    plt.tight_layout()

    # Сохраняем матрицу ошибок в файл с уникальным именем
    cm_filename = os.path.join(report_dir, f"confusion_matrix_{epoch}.png")
    plt.savefig(cm_filename)
    plt.close()

    return cm_filename



# Оценка модели с визуализацией матрицы ошибок
def evaluate_model(model, data, targets, epoch, report_dir):
    pred_labels = model.predict(data)
    accuracy = accuracy_score(targets, pred_labels)
    f1 = f1_score(targets, pred_labels, average='macro')  # Учитываем несбалансированные классы
    cm = confusion_matrix(targets, pred_labels)
    
    # Сохраняем матрицу ошибок
    cm_filename = save_confusion_matrix(cm, epoch, report_dir)

    unique, counts = np.unique(pred_labels, return_counts=True)
    class_distribution = dict(zip(unique, counts))
    return accuracy, f1, cm, cm_filename, class_distribution


# Функция для сохранения графиков Accuracy, Loss и F1 Score
def save_training_plots(accuracies, losses, f1_scores, report_dir="report_output"):
    epochs = np.arange(1, len(accuracies) + 1)

    # Accuracy Plot
    plt.figure()
    plt.plot(epochs, accuracies, label='Accuracy')
    plt.xlabel('Epochs')
    plt.ylabel('Accuracy')
    plt.title('Accuracy per Epoch')
    plt.legend()
    plt.savefig(os.path.join(report_dir, "accuracy_plot.png"), dpi=100)
    plt.close()

    # Loss Plot
    plt.figure()
    plt.plot(epochs, losses, label='Loss')
    plt.xlabel('Epochs')
    plt.ylabel('Loss')
    plt.title('Loss per Epoch')
    plt.legend()
    plt.savefig(os.path.join(report_dir, "loss_plot.png"))
    plt.close()

    # F1 Score Plot
    plt.figure()
    plt.plot(epochs, f1_scores, label='F1 Score')
    plt.xlabel('Epochs')
    plt.ylabel('F1 Score')
    plt.title('F1 Score per Epoch')
    plt.legend()
    plt.savefig(os.path.join(report_dir, "f1_score_plot.png"))
    plt.close()


def generate_reports(accuracies, losses, f1_scores, confusion_matrices_files, report_dir):
    os.makedirs(report_dir, exist_ok=True)

    # Сохранение метрик (Accuracy, Loss, F1 Score)
    save_metric_plots(accuracies, losses, f1_scores, report_dir)

    # Генерация Word-отчета
    doc = Document()
    doc.add_heading('Training Report', 0)
    
    # Добавление метрик в отчет
    doc.add_heading('Metrics Over Time', level=1)

    doc.add_picture(os.path.join(report_dir, "accuracy_plot.png"), width=Inches(5))
    doc.add_paragraph("Accuracy over epochs is shown above.")
    
    doc.add_picture(os.path.join(report_dir, "loss_plot.png"), width=Inches(5))
    doc.add_paragraph("Loss over epochs is shown above.")
    
    doc.add_picture(os.path.join(report_dir, "f1_score_plot.png"), width=Inches(5))
    doc.add_paragraph("F1 Score over epochs is shown above.")
    
    # Добавление текстового описания метрик
    add_text_summary(doc, accuracies, losses, f1_scores)

    # Добавление матриц ошибок в отчет
    for i, cm_file in enumerate(confusion_matrices_files):
        doc.add_heading(f'Confusion Matrix at Epoch {i+1}', level=2)
        doc.add_picture(cm_file, width=Inches(5))
    
    # Сохранение Word-отчета
    doc.save(os.path.join(report_dir, "report.docx"))

    # Генерация Excel-отчета с метриками
    df = pd.DataFrame({
        'Epoch': np.arange(1, len(accuracies) + 1),
        'Accuracy': accuracies,
        'Loss': losses,
        'F1 Score': f1_scores
    })
    
    df.to_excel(os.path.join(report_dir, "metrics_report.xlsx"), index=False)



def progressive_training(model, data, targets, initial_features, total_features, step=1, threshold=0.75, max_epochs=100, patience=5, learning_rate=0.05, batch_size=32, max_total_epochs=100, report_dir='report_output'):
    current_features = initial_features
    num_classes = np.unique(targets).shape[0]
    targets_one_hot = one_hot_encode(targets, num_classes)
    
    accuracies = []
    losses = []
    f1_scores = []
    confusion_matrices_files = []
    epoch_counter = 0
    total_epoch_counter = 0

    while current_features <= total_features and total_epoch_counter < max_total_epochs:
        print(f"\nTraining on {current_features} features...")

        subset_data = data[:, :current_features]

        model = ProgressiveModel(input_size=current_features, hidden_size_1=20, hidden_size_2=20, output_size=2, l2_reg=0.001, l1_reg=0.001, dropout_rate=0.2)

        for epoch in range(max_epochs):
            epoch_counter += 1
            total_epoch_counter += 1

            if total_epoch_counter >= max_total_epochs:
                print(f"Reached the maximum total epochs ({max_total_epochs}). Stopping training.")
                break

            total_loss = 0

            for mini_batch_data, mini_batch_targets in create_mini_batches(subset_data, targets_one_hot, batch_size):
                predictions = model.forward(mini_batch_data, train=True)
                model.backward(mini_batch_data, mini_batch_targets, learning_rate=learning_rate)
                
                batch_loss = -np.sum(mini_batch_targets * np.log(predictions)) / mini_batch_data.shape[0]
                total_loss += batch_loss

            # Оценка модели с сохранением матрицы ошибок
            accuracy, f1, cm, cm_filename, class_distribution = evaluate_model(model, subset_data, targets, epoch_counter, report_dir)

            accuracies.append(accuracy)
            losses.append(total_loss / (subset_data.shape[0] / batch_size))
            f1_scores.append(f1)
            confusion_matrices_files.append(cm_filename)

            print(f"Epoch {epoch_counter}: Accuracy = {accuracy:.2f}, Loss = {total_loss:.4f}, F1 Score = {f1:.2f}")
            print(f"Predicted class distribution: {class_distribution}")

            if early_stopping_check(accuracies, patience=patience):
                print(f"Early stopping at epoch {epoch_counter}")
                break

            if accuracy >= threshold:
                print(f"Model reached the accuracy threshold at {current_features} features. Adding more features.")
                current_features += step
                break

            learning_rate = adjust_learning_rate(learning_rate, epoch_counter)
        
        data = generate_new_features(data, max_new_features=100)
        data = augment_data(data, noise_level=0.05)

        if current_features > total_features:
            print("Progressive training completed.")
            break

    # Генерация отчета после обучения
    generate_reports(accuracies, losses, f1_scores, confusion_matrices_files, report_dir)
